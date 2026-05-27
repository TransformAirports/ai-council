"""Build the one-page IT-department memo describing the AI Council system.

Run:  .venv/bin/python scripts/build_it_memo.py
Out:  it-department-memo.docx (at the repo root)
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parent.parent
OUT_PATH = REPO_ROOT / "it-department-memo.docx"


def _set_margins(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)


def _set_default_font(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)


def _title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Calibri"


def _header_row(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(f"{label:<7}")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.name = "Calibri"
    r2 = p.add_run(value)
    r2.font.size = Pt(10)
    r2.font.name = "Calibri"


def _rule(doc: Document) -> None:
    """Thin horizontal rule via an underlined empty run."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "888888")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _para(doc: Document, body: str, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        r.font.size = Pt(11)
        r.font.name = "Calibri"
        r2 = p.add_run(" " + body)
        r2.font.size = Pt(11)
        r2.font.name = "Calibri"
    else:
        r = p.add_run(body)
        r.font.size = Pt(11)
        r.font.name = "Calibri"


def build(out_path: Path = OUT_PATH) -> Path:
    doc = Document()
    _set_margins(doc)
    _set_default_font(doc)

    _title(doc, "MEMORANDUM")
    _header_row(doc, "TO:", "Information Technology")
    _header_row(doc, "FROM:", "Christian Kessler IV")
    _header_row(doc, "DATE:", date.today().strftime("%B %d, %Y"))
    _header_row(
        doc,
        "RE:",
        "Production multi-agent AI system — architecture and operational disciplines",
    )
    _rule(doc)

    _para(
        doc,
        "This memo documents the multi-agent AI orchestration system I operate "
        "for executive-level analytical work. It is provided so the department has "
        "visibility into the stack, the controls in place, and the engineering "
        "practices that govern its use.",
    )

    _para(
        doc,
        "A nineteen-agent system — fifteen research lenses plus four process "
        "agents — built on the Claude Agent SDK (Python 3.11+, asyncio), driving "
        "Claude Opus and Sonnet in role-specific allocations. Agents are defined "
        "as versioned markdown files with YAML frontmatter under "
        "`.claude/agents/`; the orchestrator parses tool allowlists, model "
        "selection, and system prompts at runtime. Tool permissions are scoped "
        "per-agent (WebSearch, WebFetch, Read, Write); `permission_mode` is set "
        "to `bypassPermissions` only inside the orchestrator's controlled "
        "subprocess scope, never broader.",
        bold_lead="Stack.",
    )

    _para(
        doc,
        "Four stages with strict separation. (1) Parallel independent research "
        "via `asyncio.gather` — agents are explicitly prohibited from reading "
        "peers' output to force evidence variance rather than convergence. "
        "(2) Sequential adversarial synthesis with Strategist–Red Team "
        "alternation across three revision rounds, executed as a state machine, "
        "not a conversation. (3) Edit and source-verification, with the "
        "Fact-checker holding veto over any claim that does not trace to a "
        "Stage 1 brief. (4) Deterministic DOCX generation through a custom "
        "`python-docx` builder. Two formal human checkpoints separate Stages "
        "2–3 and 3–4.",
        bold_lead="Pipeline.",
    )

    _para(
        doc,
        "Per-step cost tallying and retrospective generation, with full "
        "intermediate-stage preservation under date-slugged immutable archives. "
        "On-disk state inspection enabling mid-pipeline crash recovery via a "
        "`--resume` mode that re-runs only the artifacts missing from the "
        "previous attempt. `max_turns` budgets tuned to file-read counts to "
        "prevent a specific Claude Code failure mode in which the subprocess "
        "emits `is_error=true / subtype=success` envelopes when an agent "
        "exhausts its turn budget before writing — diagnosed by inspecting the "
        "SDK's result-message reducer, mitigated by a post-call write-verifier "
        "that surfaces silent failures with an actionable error. Authentication "
        "routes through either a Claude.ai subscription or an Anthropic API "
        "key; the orchestrator inherits whichever is configured in the "
        "underlying CLI binary.",
        bold_lead="Operational disciplines.",
    )

    _para(
        doc,
        "The non-obvious engineering is not in writing the agents — it is in "
        "tuning how they interact. Calibrating the Strategist to absorb "
        "adversarial critique without losing narrative voice. Constraining the "
        "Red Team to attack content rather than style. Sequencing revision "
        "rounds so the argument converges without flattening. Composing the "
        "per-thesis council from twelve domain-specialized lenses plus two "
        "deliberately unstructured ones — an unprepared-by-design “Slacker” "
        "and an operator-modeled “Virtual Christian” — to inject failure "
        "modes the structured agents will not produce on their own. This is "
        "system design, not prompt engineering.",
        bold_lead="Where the engineering actually lives.",
    )

    _para(
        doc,
        "Agent definitions are modified only through reviewed pull requests. No "
        "agent edits itself or another. Every run is auditable end-to-end — "
        "research briefs, draft history, critique rounds, fact-check report, "
        "and final document are preserved together.",
    )

    _para(
        doc,
        "I am available to walk anyone in the department through the codebase, "
        "the orchestrator, or the operational logs.",
    )

    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    path = build()
    print(f"Wrote: {path.relative_to(REPO_ROOT)}")
