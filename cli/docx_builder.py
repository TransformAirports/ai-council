"""Stage 4 — produce airport-grade Word reports and decision briefs.

Replicates what the Claude Code `docx` skill would otherwise do, but as plain
Python so the CLI does not depend on Claude Code being present at runtime.

The converter preserves the structures Council reports actually use:
headings, paragraphs, lists, blockquotes, pipe tables, inline emphasis, and
numbered notes.  It also builds a deterministic executive decision brief from
the verified argument and runs a structural publishing gate before returning.
"""
from __future__ import annotations

import json
import re
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from urllib.parse import urlsplit

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from cli.publishing_quality import (
    QualityIssue,
    QualityReport,
    assert_quality,
    lint_markdown,
    prepare_word_visual_inspection_receipt,
    qa_docx,
    render_office_artifact,
)


HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
BULLET_RE = re.compile(r"^[\-\*]\s+(.*)$")
ORDERED_RE = re.compile(r"^\d+\.\s+(.*)$")
BLOCKQUOTE_RE = re.compile(r"^>\s?(.*)$")
INLINE_BOLD = re.compile(r"\*\*(.+?)\*\*")
INLINE_ITALIC = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
INLINE_CODE = re.compile(r"`([^`]+)`")
HTML_COMMENT = re.compile(r"<!--.*?-->", re.DOTALL)
PIPE_TABLE_RE = re.compile(r"^\s*\|?.+\|.+\|?\s*$")
PIPE_DIVIDER_RE = re.compile(
    r"^\s*\|?\s*:?-{3,}:?\s*(?:\|\s*:?-{3,}:?\s*)+\|?\s*$"
)
HORIZONTAL_RULE_RE = re.compile(r"^\s*(?:-{3,}|\*{3,}|_{3,})\s*$")

# Markdown footnotes: [^4] markers in text, "[^4]: source" definitions at the
# document end. Rendered as superscript numbers plus a styled Notes section.
FOOTNOTE_DEF_RE = re.compile(r"^\[\^(\d+)\]:\s*(.*)$")
FOOTNOTE_MARK_RE = re.compile(r"\[\^(\d+)\]")
SOURCE_URL_RE = re.compile(r"https?://[^\s]+")

# Legacy internal-provenance tags from older runs ("[Economist brief,
# Finding 3]"). These name the Council's machinery, not a source — strip them
# from anything reader-facing.
INTERNAL_CITATION_RE = re.compile(
    r"\s?\[[^\]]*\b(?:brief|Stage\s*1)\b[^\]]*\]", re.IGNORECASE
)
INTERNAL_STAGE_LINE_RE = re.compile(
    r"(?im)^\s*#{0,6}\s*.*\bStage\s+[123]\b.*"
    r"\b(?:draft|fact[- ]?checked|output)\b.*(?:\n|$)"
)
LEGACY_SOURCE_TAG_RE = re.compile(r"\s*\[Source:\s*[^\]]+\]", re.IGNORECASE)


# Shared airport-grade visual system. The JSON in assets/brand is the canonical
# machine-readable contract; these constants keep document production usable
# when the repository is packaged without non-Python data files.
DISPLAY_FONT = "Georgia"
BODY_FONT = "Aptos"
BODY_FONT_FALLBACK = "Calibri"
RUNWAY_NAVY = RGBColor(0x0B, 0x2D, 0x4D)
TERMINAL_BLUE = RGBColor(0x2E, 0x84, 0xA5)
GUIDANCE_GOLD = RGBColor(0xD4, 0xA2, 0x4C)
OPERATIONS_SLATE = RGBColor(0x41, 0x56, 0x69)
APRON_FOG = "EDF3F6"
SIGNAL_GREEN = RGBColor(0x24, 0x74, 0x5C)
ALERT_RED = RGBColor(0xA6, 0x41, 0x3A)
INK = RGBColor(0x17, 0x23, 0x2D)
REPO_ROOT = Path(__file__).resolve().parent.parent
COUNCIL_LOGO = REPO_ROOT / "assets" / "council-logo.png"
AI_ACCOUNTABILITY_NOTICE = (
    "This decision-support document was generated with assistance from a "
    "multi-model AI research system. A named human decision owner remains "
    "responsible for verifying the evidence, judging local applicability, "
    "and approving any action."
)


def strip_internal_citations(text: str) -> str:
    """Remove agent/brief provenance tags that should never reach a reader."""
    return INTERNAL_CITATION_RE.sub("", text)


def sanitize_reader_markdown(text: str) -> str:
    """Remove legacy production annotations from reader-facing Markdown.

    Legacy ``[Source: ...]`` tags cannot be converted into trustworthy
    footnotes because they do not carry a stable claim/source relationship.
    New runs use numbered notes; older internal tags are stripped and the
    publishing linter prevents any unrecognized residue from passing.
    """
    text = _strip_comments(text)
    text = INTERNAL_STAGE_LINE_RE.sub("", text)
    text = strip_internal_citations(text)
    text = LEGACY_SOURCE_TAG_RE.sub("", text)
    return re.sub(r"\n{3,}", "\n\n", text).strip() + "\n"


def _compact_argument_memo_source_urls(markdown: str) -> str:
    """Keep one-page memo citations traceable without printing long raw URLs."""

    def compact(match: re.Match[str]) -> str:
        raw = match.group(0)
        clean = raw.rstrip(".,;")
        trailing = raw[len(clean):]
        host = urlsplit(clean).netloc.casefold()
        if host.startswith("www."):
            host = host[4:]
        return (host or clean) + trailing

    lines: list[str] = []
    for line in markdown.splitlines():
        if FOOTNOTE_DEF_RE.match(line.strip()):
            line = SOURCE_URL_RE.sub(compact, line)
        lines.append(line)
    return "\n".join(lines)


def _add_inline(paragraph, text: str, base_size: int = 11, font: str = "Calibri") -> None:
    """Add a text run, parsing **bold**, *italic*, `code`, and [^n] footnotes."""
    if not text:
        return
    pos = 0
    tokens: list[tuple[str, str]] = []
    pattern = re.compile(
        r"\*\*(.+?)\*\*|(?<!\*)\*([^*]+?)\*(?!\*)|`([^`]+)`|\[\^(\d+)\]"
    )
    for m in pattern.finditer(text):
        if m.start() > pos:
            tokens.append(("plain", text[pos : m.start()]))
        if m.group(1) is not None:
            tokens.append(("bold", m.group(1)))
        elif m.group(2) is not None:
            tokens.append(("italic", m.group(2)))
        elif m.group(3) is not None:
            tokens.append(("code", m.group(3)))
        elif m.group(4) is not None:
            tokens.append(("footnote", m.group(4)))
        pos = m.end()
    if pos < len(text):
        tokens.append(("plain", text[pos:]))

    for kind, value in tokens:
        run = paragraph.add_run(value)
        run.font.size = Pt(base_size)
        run.font.name = font
        if kind == "bold":
            run.bold = True
        elif kind == "italic":
            run.italic = True
        elif kind == "code":
            run.font.name = "Consolas"
        elif kind == "footnote":
            run.font.superscript = True
            run.font.size = Pt(max(base_size - 2, 7))


def _strip_comments(text: str) -> str:
    return HTML_COMMENT.sub("", text)


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell, *, top: int = 90, start: int = 110,
                      bottom: int = 90, end: int = 110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _split_pipe_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [cell.strip() for cell in stripped.split("|")]


def _add_markdown_table(
    doc: Document,
    header: list[str],
    rows: list[list[str]],
    *,
    body_size: float,
    font: str,
) -> None:
    columns = len(header)
    table = doc.add_table(rows=1, cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    try:
        table.style = "Light Shading Accent 1"
    except KeyError:
        table.style = "Table Grid"
    _set_repeat_table_header(table.rows[0])

    for index, value in enumerate(header):
        cell = table.rows[0].cells[index]
        _set_cell_shading(cell, "0B2D4D")
        _set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        _add_inline(paragraph, value, base_size=max(body_size - 0.5, 8.5), font=font)
        for run in paragraph.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for row_index, values in enumerate(rows, 1):
        row = table.add_row()
        values = values[:columns] + [""] * max(0, columns - len(values))
        for column, value in enumerate(values):
            cell = row.cells[column]
            _set_cell_margins(cell)
            if row_index % 2 == 0:
                _set_cell_shading(cell, APRON_FOG)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            _add_inline(paragraph, value, base_size=max(body_size - 0.5, 8.5), font=font)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def _style_callout(paragraph) -> None:
    """Apply a restrained airport-blue callout treatment to blockquotes."""
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), APRON_FOG)
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "2E84A5")
    borders.append(left)
    p_pr.append(borders)
    paragraph.paragraph_format.left_indent = Pt(16)
    paragraph.paragraph_format.right_indent = Pt(8)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(8)


def _add_horizontal_rule(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "7")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D4A24C")
    borders.append(bottom)
    p_pr.append(borders)


def _markdown_to_docx(doc: Document, markdown: str, body_size: int = 11, font: str = "Calibri") -> None:
    text = sanitize_reader_markdown(markdown)
    # Pull footnote definitions out of the body; render them as a Notes
    # section after the content, in the order the author numbered them.
    notes: list[tuple[str, str]] = []
    body_lines: list[str] = []
    for raw in text.splitlines():
        m = FOOTNOTE_DEF_RE.match(raw.strip())
        if m:
            notes.append((m.group(1), m.group(2)))
        else:
            body_lines.append(raw)
    lines = body_lines
    i = 0
    while i < len(lines):
        raw = lines[i].rstrip()
        if not raw.strip():
            i += 1
            continue
        if HORIZONTAL_RULE_RE.match(raw):
            _add_horizontal_rule(doc)
            i += 1
            continue

        m = HEADING_RE.match(raw)
        if m:
            level = min(len(m.group(1)), 4)
            heading = doc.add_heading(level=level)
            heading_sizes = {1: 19, 2: 14, 3: 12, 4: 10.5}
            _add_inline(
                heading,
                m.group(2),
                base_size=heading_sizes[level],
                font=DISPLAY_FONT,
            )
            i += 1
            continue

        # GitHub-style pipe tables. A divider row immediately after the
        # header distinguishes a table from prose containing a vertical bar.
        if (
            PIPE_TABLE_RE.match(raw)
            and i + 1 < len(lines)
            and PIPE_DIVIDER_RE.match(lines[i + 1].rstrip())
        ):
            header = _split_pipe_row(raw)
            i += 2
            rows: list[list[str]] = []
            while i < len(lines) and PIPE_TABLE_RE.match(lines[i].rstrip()):
                rows.append(_split_pipe_row(lines[i].rstrip()))
                i += 1
            _add_markdown_table(
                doc,
                header,
                rows,
                body_size=body_size,
                font=font,
            )
            continue

        m = BULLET_RE.match(raw)
        if m:
            while i < len(lines):
                line = lines[i].rstrip()
                bm = BULLET_RE.match(line)
                if not bm:
                    break
                p = doc.add_paragraph(style="List Bullet")
                _add_inline(p, bm.group(1), base_size=body_size, font=font)
                i += 1
            continue

        m = ORDERED_RE.match(raw)
        if m:
            while i < len(lines):
                line = lines[i].rstrip()
                om = ORDERED_RE.match(line)
                if not om:
                    break
                p = doc.add_paragraph(style="List Number")
                _add_inline(p, om.group(1), base_size=body_size, font=font)
                i += 1
            continue

        m = BLOCKQUOTE_RE.match(raw)
        if m:
            buf: list[str] = []
            while i < len(lines):
                line = lines[i].rstrip()
                bm = BLOCKQUOTE_RE.match(line)
                if not bm:
                    break
                buf.append(bm.group(1))
                i += 1
            p = doc.add_paragraph()
            _add_inline(p, " ".join(buf), base_size=body_size, font=font)
            for run in p.runs:
                run.font.color.rgb = RUNWAY_NAVY
            _style_callout(p)
            continue

        # Plain paragraph: collect until blank line.
        buf = [raw]
        i += 1
        while i < len(lines) and lines[i].strip():
            if HEADING_RE.match(lines[i]) or BULLET_RE.match(lines[i]) or ORDERED_RE.match(lines[i]):
                break
            buf.append(lines[i].rstrip())
            i += 1
        p = doc.add_paragraph()
        _add_inline(p, " ".join(buf), base_size=body_size, font=font)
        p.paragraph_format.line_spacing = 1.16
        p.paragraph_format.space_after = Pt(7)

    # The Notes section — the document's citations, styled quietly.
    if notes:
        heading = doc.add_heading(level=2)
        _add_inline(heading, "Notes", base_size=14, font=DISPLAY_FONT)
        for num, note_text in notes:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            mark = p.add_run(num)
            mark.font.superscript = True
            mark.font.size = Pt(8)
            mark.font.name = font
            p.add_run("  ")
            _add_inline(p, note_text, base_size=max(body_size - 1.5, 8), font=font)
            for run in p.runs[2:]:
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def _set_default_font(doc: Document, name: str = BODY_FONT, size: float = 10.5) -> None:
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = INK
    style.paragraph_format.space_after = Pt(7)
    style.paragraph_format.line_spacing = 1.16


def _add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, end))


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def _configure_document(doc: Document, *, compact: bool = False) -> None:
    """Apply the Council's Word grid, hierarchy, header, and footer."""
    body_size = 10 if compact else 10.5
    _set_default_font(doc, BODY_FONT, body_size)
    section = doc.sections[0]
    section.top_margin = Inches(0.68)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    for level, size, before, after in (
        (1, 19, 16, 8),
        (2, 14, 13, 5),
        (3, 12, 10, 3),
        (4, 10.5, 8, 2),
    ):
        style = doc.styles[f"Heading {level}"]
        style.font.name = DISPLAY_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RUNWAY_NAVY
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style in ("List Bullet", "List Number"):
        try:
            style = doc.styles[list_style]
        except KeyError:
            continue
        style.font.name = BODY_FONT
        style.font.size = Pt(body_size)
        style.paragraph_format.space_after = Pt(4)

    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("TRANSFORM AIRPORTS  /  AI RESEARCH COUNCIL")
    r.font.name = BODY_FONT
    r.font.size = Pt(7.5)
    r.font.bold = True
    r.font.color.rgb = OPERATIONS_SLATE

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("CONFIDENTIAL WORKING PAPER  •  ")
    r.font.name = BODY_FONT
    r.font.size = Pt(7.5)
    r.font.color.rgb = OPERATIONS_SLATE
    _add_page_field(p)

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def _add_cover_page(
    doc: Document,
    title: str,
    subtitle: str,
    revision_label: str | None = None,
) -> None:
    # Editorial, left-aligned cover: the title gets width and avoids the
    # awkward center-wrapped stack produced by the original template.
    if COUNCIL_LOGO.is_file():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run().add_picture(str(COUNCIL_LOGO), width=Inches(2.75))
    for _ in range(2):
        doc.add_paragraph()

    eyebrow = doc.add_paragraph()
    eyebrow.paragraph_format.space_after = Pt(8)
    run = eyebrow.add_run("AIRPORT EXECUTIVE ANALYSIS")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = BODY_FONT
    run.font.color.rgb = TERMINAL_BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(30)
    run.font.name = DISPLAY_FONT
    run.font.color.rgb = RUNWAY_NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(subtitle)
    run.font.size = Pt(14)
    run.font.name = BODY_FONT
    run.font.color.rgb = OPERATIONS_SLATE

    if revision_label:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(8)
        run = p.add_run(revision_label.upper())
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = BODY_FONT
        run.font.color.rgb = GUIDANCE_GOLD

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Transform Airports AI Council")
    run.font.size = Pt(12)
    run.font.name = BODY_FONT
    run.bold = True
    run.font.color.rgb = RUNWAY_NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(date.today().strftime("%B %d, %Y"))
    run.font.size = Pt(11)
    run.font.name = BODY_FONT
    run.font.color.rgb = OPERATIONS_SLATE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("DECISION SUPPORT  •  HUMAN REVIEW REQUIRED")
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = BODY_FONT
    run.font.color.rgb = ALERT_RED

    _add_ai_accountability_notice(doc)
    doc.add_page_break()


def _add_ai_accountability_notice(
    doc: Document,
    *,
    compact: bool = False,
) -> None:
    """State the AI/human boundary plainly in every released Word artifact."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Pt(10)
    p.paragraph_format.right_indent = Pt(10)
    _add_inline(
        p,
        AI_ACCOUNTABILITY_NOTICE,
        base_size=8.5 if compact else 9,
        font=BODY_FONT,
    )
    for run in p.runs:
        run.font.color.rgb = OPERATIONS_SLATE
    _style_callout(p)


def _add_table_of_contents(doc: Document) -> None:
    heading = doc.add_heading("Contents", level=1)
    heading.paragraph_format.space_after = Pt(12)
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = run._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
    run._element.append(fld_begin)
    instr = run._element.makeelement(qn("w:instrText"), {qn("xml:space"): "preserve"})
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._element.append(instr)
    fld_sep = run._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "separate"})
    run._element.append(fld_sep)
    fld_end = run._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
    fld_end.set(qn("w:dirty"), "true")
    run._element.append(fld_end)
    doc.add_page_break()


def _build_full_report(
    title: str,
    final_draft_md: str,
    methodology_md: str,
    out_path: Path,
    visual_brief: dict | None = None,
    decision_context: DecisionContext | dict | None = None,
    revision_label: str | None = None,
) -> None:
    assert_quality(lint_markdown(final_draft_md, location="final draft"))
    doc = Document()
    _configure_document(doc)
    _add_cover_page(
        doc,
        title=title,
        subtitle="Independent analysis for airport decision-makers",
        revision_label=revision_label,
    )
    brief = _build_decision_brief(
        final_draft_md,
        visual_brief=visual_brief,
        decision_context=decision_context,
    )
    _add_decision_brief(doc, brief)
    doc.add_page_break()
    _add_table_of_contents(doc)
    _markdown_to_docx(doc, final_draft_md, body_size=11)
    if _renderable_exhibits(visual_brief):
        doc.add_page_break()
        _add_decision_exhibits(doc, visual_brief)
    doc.add_page_break()
    _add_technical_evidence_appendix(doc, final_draft_md, visual_brief)
    doc.add_page_break()
    doc.add_heading("Technical appendix: Methodology", level=1)
    _markdown_to_docx(doc, methodology_md, body_size=11)
    doc.save(out_path)


@dataclass(frozen=True)
class _ContentUnit:
    section: str
    kind: str
    text: str
    index: int


@dataclass(frozen=True)
class DecisionBrief:
    bottom_line: str
    why_now: tuple[str, ...]
    evidence: tuple[str, ...]
    recommendations: tuple[str, ...]
    risks: tuple[str, ...]
    notes: tuple[tuple[str, str], ...] = ()
    decision_owner: str = "Assign the accountable decision owner before authorization."
    approval_route: str = "Confirm the governing approval route before authorization."
    first_90_day_action: str = (
        "Define and authorize the first 90-day work plan before launch."
    )
    success_measures: tuple[str, ...] = (
        "Set measurable acceptance, stop, and reporting criteria before launch.",
    )
    time_horizon: str = ""


@dataclass(frozen=True)
class DecisionContext:
    """Structured run-prompt decision fields carried into Word production."""

    decision: str = ""
    decision_owner: str = ""
    approval_path: str = ""
    first_action: str = ""
    time_horizon: str = ""
    success_measures: tuple[str, ...] = ()


@dataclass(frozen=True)
class _ExhibitColumn:
    label: str
    unit: str = ""


@dataclass(frozen=True)
class _ExhibitRow:
    label: str
    values: tuple[str, ...]
    unit: str = ""


@dataclass(frozen=True)
class _ExhibitStep:
    label: str
    detail: str
    timing: str = ""
    owner: str = ""
    trigger: str = ""
    success_measure: str = ""


@dataclass(frozen=True)
class _RenderableExhibit:
    title: str
    exhibit_type: str
    takeaway: str
    evidence_ids: tuple[str, ...]
    source_note: str
    row_header: str = ""
    columns: tuple[_ExhibitColumn, ...] = ()
    rows: tuple[_ExhibitRow, ...] = ()
    steps: tuple[_ExhibitStep, ...] = ()


def _clean_text(value: object) -> str:
    if not isinstance(value, str):
        return ""
    return " ".join(value.split()).strip()


def _string_values(value: object) -> tuple[str, ...]:
    """Normalize permissive visual-brief values without inventing content."""
    if isinstance(value, str):
        cleaned = _clean_text(value)
        return (cleaned,) if cleaned else ()
    if isinstance(value, (list, tuple)):
        values: list[str] = []
        for item in value:
            values.extend(_string_values(item))
        return tuple(dict.fromkeys(values))
    if isinstance(value, dict):
        values: list[str] = []
        for item in value.values():
            values.extend(_string_values(item))
        return tuple(dict.fromkeys(values))
    return ()


def _coerce_decision_context(
    value: DecisionContext | dict | None,
) -> DecisionContext:
    """Normalize direct run metadata without inferring fields from prose."""

    if isinstance(value, DecisionContext):
        return value
    if not isinstance(value, dict):
        return DecisionContext()

    approval_values = _string_values(
        value.get("approval_path") or value.get("approval_route")
    )
    success_measures = _string_values(
        value.get("success_measures") or value.get("success_measure")
    )
    return DecisionContext(
        decision=_clean_text(value.get("decision") or value.get("decision_required")),
        decision_owner=_clean_text(value.get("decision_owner")),
        approval_path=(
            " → ".join(approval_values)
            if len(approval_values) > 1
            else (approval_values[0] if approval_values else "")
        ),
        first_action=_clean_text(
            value.get("first_action") or value.get("first_90_day_action")
        ),
        time_horizon=_clean_text(value.get("time_horizon")),
        success_measures=success_measures,
    )


def _apply_structured_decision_context(
    brief: DecisionBrief,
    decision_context: DecisionContext | dict | None,
) -> DecisionBrief:
    """Overlay authoritative run fields after any Art Director enrichment."""

    context = _coerce_decision_context(decision_context)
    if not any(
        (
            context.decision,
            context.decision_owner,
            context.approval_path,
            context.first_action,
            context.time_horizon,
            context.success_measures,
        )
    ):
        return brief
    return DecisionBrief(
        bottom_line=context.decision or brief.bottom_line,
        why_now=brief.why_now,
        evidence=brief.evidence,
        recommendations=brief.recommendations,
        risks=brief.risks,
        notes=brief.notes,
        decision_owner=context.decision_owner or brief.decision_owner,
        approval_route=context.approval_path or brief.approval_route,
        first_90_day_action=context.first_action or brief.first_90_day_action,
        success_measures=context.success_measures or brief.success_measures,
        time_horizon=context.time_horizon or brief.time_horizon,
    )


def _visual_brief_containers(visual_brief: dict | None) -> tuple[dict, ...]:
    if not isinstance(visual_brief, dict):
        return ()
    containers = [visual_brief]
    for key in (
        "decision_frame",
        "decision_metadata",
        "governance",
        "implementation",
        "execution",
    ):
        nested = visual_brief.get(key)
        if isinstance(nested, dict):
            containers.append(nested)
    return tuple(containers)


def _visual_brief_value(
    visual_brief: dict | None,
    *keys: str,
) -> object | None:
    for container in _visual_brief_containers(visual_brief):
        for key in keys:
            value = container.get(key)
            if value not in (None, "", [], {}):
                return value
    return None


def _apply_visual_decision_context(
    brief: DecisionBrief,
    visual_brief: dict | None,
) -> DecisionBrief:
    """Overlay Art Director decision metadata, retaining explicit fallbacks."""
    if not visual_brief:
        return brief

    decision = _clean_text(
        _visual_brief_value(visual_brief, "decision") or brief.bottom_line
    )
    owner = _clean_text(
        _visual_brief_value(
            visual_brief,
            "decision_owner",
            "owner",
            "accountable_owner",
            "accountable_executive",
        )
        or brief.decision_owner
    )
    approval_values = _string_values(
        _visual_brief_value(
            visual_brief,
            "approval_route",
            "approval_path",
            "approvals",
        )
    )
    approval_route = (
        " → ".join(approval_values)
        if len(approval_values) > 1
        else (approval_values[0] if approval_values else brief.approval_route)
    )
    first_action_values = _string_values(
        _visual_brief_value(
            visual_brief,
            "first_90_day_action",
            "first_90_days",
            "first_90_day_move",
            "initial_action",
        )
    )
    first_action = (
        " ".join(first_action_values)
        if first_action_values
        else brief.first_90_day_action
    )
    success_measures = _string_values(
        _visual_brief_value(
            visual_brief,
            "success_measures",
            "success_measure",
            "measures_of_success",
            "success_criteria",
        )
    ) or brief.success_measures

    return DecisionBrief(
        bottom_line=decision or brief.bottom_line,
        why_now=brief.why_now,
        evidence=brief.evidence,
        recommendations=brief.recommendations,
        risks=brief.risks,
        notes=brief.notes,
        decision_owner=owner,
        approval_route=approval_route,
        first_90_day_action=first_action,
        success_measures=success_measures,
        time_horizon=brief.time_horizon,
    )


def _required_exhibit_text(payload: dict, key: str, *, exhibit: int) -> str:
    value = _clean_text(payload.get(key))
    if not value:
        raise ValueError(
            f"report_visuals[{exhibit}] requires a non-empty {key!r}"
        )
    return value


def _exhibit_cell_text(
    value: object,
    *,
    exhibit: int,
    row: int,
    column: int,
    unit_available: bool,
) -> str:
    """Preserve supplied table values without deriving or completing them."""
    if isinstance(value, bool) or not isinstance(value, (str, int, float)):
        raise ValueError(
            "report_visuals"
            f"[{exhibit}].rows[{row}].values[{column}] must be text or a number"
        )
    if isinstance(value, str):
        cleaned = _clean_text(value)
        if not cleaned:
            raise ValueError(
                "report_visuals"
                f"[{exhibit}].rows[{row}].values[{column}] cannot be blank"
            )
        return cleaned
    if not unit_available:
        raise ValueError(
            "report_visuals"
            f"[{exhibit}].rows[{row}].values[{column}] is numeric but has no "
            "row or column unit"
        )
    return str(value)


def _renderable_exhibits(
    visual_brief: dict | None,
) -> tuple[_RenderableExhibit, ...]:
    """Parse only complete exhibit data the Word builder can actually render.

    The signature visual and slide-level map/chart instructions remain internal
    production inputs. Legacy report entries without a supported
    ``exhibit_type`` are deliberately ignored instead of being relabeled as
    finished exhibits. Recognized structures fail closed when required source
    or data fields are missing, preventing the builder from filling gaps.
    """
    if not isinstance(visual_brief, dict):
        return ()
    payloads = visual_brief.get("report_visuals")
    if not isinstance(payloads, list):
        return ()

    exhibits: list[_RenderableExhibit] = []
    supported = {"table", "comparison", "flow", "timeline"}
    for exhibit_index, payload in enumerate(payloads):
        if not isinstance(payload, dict):
            continue
        exhibit_type = _clean_text(payload.get("exhibit_type")).lower()
        if exhibit_type not in supported:
            continue

        title = _required_exhibit_text(
            payload, "title", exhibit=exhibit_index
        )
        takeaway = _required_exhibit_text(
            payload, "takeaway", exhibit=exhibit_index
        )
        source_note = _required_exhibit_text(
            payload, "source_note", exhibit=exhibit_index
        )
        evidence_ids = _string_values(payload.get("evidence_ids"))
        if not evidence_ids:
            raise ValueError(
                f"report_visuals[{exhibit_index}] requires evidence_ids"
            )

        if exhibit_type in {"table", "comparison"}:
            row_header = _required_exhibit_text(
                payload, "row_header", exhibit=exhibit_index
            )
            raw_columns = payload.get("columns")
            raw_rows = payload.get("rows")
            if not isinstance(raw_columns, list) or not raw_columns:
                raise ValueError(
                    f"report_visuals[{exhibit_index}] requires columns"
                )
            if not isinstance(raw_rows, list) or not raw_rows:
                raise ValueError(
                    f"report_visuals[{exhibit_index}] requires rows"
                )
            columns: list[_ExhibitColumn] = []
            for column_index, column in enumerate(raw_columns):
                if not isinstance(column, dict):
                    raise ValueError(
                        "report_visuals"
                        f"[{exhibit_index}].columns[{column_index}] must be an object"
                    )
                label = _clean_text(column.get("label"))
                if not label:
                    raise ValueError(
                        "report_visuals"
                        f"[{exhibit_index}].columns[{column_index}] requires label"
                    )
                columns.append(
                    _ExhibitColumn(
                        label=label,
                        unit=_clean_text(column.get("unit")),
                    )
                )
            rows: list[_ExhibitRow] = []
            for row_index, row in enumerate(raw_rows):
                if not isinstance(row, dict):
                    raise ValueError(
                        "report_visuals"
                        f"[{exhibit_index}].rows[{row_index}] must be an object"
                    )
                label = _clean_text(row.get("label"))
                raw_values = row.get("values")
                row_unit = _clean_text(row.get("unit"))
                if not label or not isinstance(raw_values, list):
                    raise ValueError(
                        "report_visuals"
                        f"[{exhibit_index}].rows[{row_index}] requires label and values"
                    )
                if len(raw_values) != len(columns):
                    raise ValueError(
                        "report_visuals"
                        f"[{exhibit_index}].rows[{row_index}] has "
                        f"{len(raw_values)} values for {len(columns)} columns"
                    )
                values = tuple(
                    _exhibit_cell_text(
                        value,
                        exhibit=exhibit_index,
                        row=row_index,
                        column=column_index,
                        unit_available=bool(row_unit or columns[column_index].unit),
                    )
                    for column_index, value in enumerate(raw_values)
                )
                rows.append(_ExhibitRow(label=label, values=values, unit=row_unit))
            exhibits.append(
                _RenderableExhibit(
                    title=title,
                    exhibit_type=exhibit_type,
                    takeaway=takeaway,
                    evidence_ids=evidence_ids,
                    source_note=source_note,
                    row_header=row_header,
                    columns=tuple(columns),
                    rows=tuple(rows),
                )
            )
            continue

        if exhibit_type == "flow":
            raw_steps = payload.get("steps")
            if not isinstance(raw_steps, list) or len(raw_steps) < 2:
                raise ValueError(
                    f"report_visuals[{exhibit_index}] requires at least two steps"
                )
            steps = []
            for step_index, step in enumerate(raw_steps):
                if not isinstance(step, dict):
                    raise ValueError(
                        "report_visuals"
                        f"[{exhibit_index}].steps[{step_index}] must be an object"
                    )
                label = _clean_text(step.get("label"))
                detail = _clean_text(step.get("detail"))
                if not label or not detail:
                    raise ValueError(
                        "report_visuals"
                        f"[{exhibit_index}].steps[{step_index}] requires label and detail"
                    )
                steps.append(
                    _ExhibitStep(
                        label=label,
                        detail=detail,
                        owner=_clean_text(step.get("owner")),
                        trigger=_clean_text(step.get("trigger")),
                    )
                )
            exhibits.append(
                _RenderableExhibit(
                    title=title,
                    exhibit_type=exhibit_type,
                    takeaway=takeaway,
                    evidence_ids=evidence_ids,
                    source_note=source_note,
                    steps=tuple(steps),
                )
            )
            continue

        raw_milestones = payload.get("milestones")
        if not isinstance(raw_milestones, list) or len(raw_milestones) < 2:
            raise ValueError(
                f"report_visuals[{exhibit_index}] requires at least two milestones"
            )
        milestones = []
        for milestone_index, milestone in enumerate(raw_milestones):
            if not isinstance(milestone, dict):
                raise ValueError(
                    "report_visuals"
                    f"[{exhibit_index}].milestones[{milestone_index}] must be an object"
                )
            period = _clean_text(milestone.get("period"))
            action = _clean_text(milestone.get("action"))
            if not period or not action:
                raise ValueError(
                    "report_visuals"
                    f"[{exhibit_index}].milestones[{milestone_index}] "
                    "requires period and action"
                )
            milestones.append(
                _ExhibitStep(
                    label=action,
                    detail=_clean_text(milestone.get("detail")),
                    timing=period,
                    owner=_clean_text(milestone.get("owner")),
                    success_measure=_clean_text(
                        milestone.get("success_measure")
                    ),
                )
            )
        exhibits.append(
            _RenderableExhibit(
                title=title,
                exhibit_type=exhibit_type,
                takeaway=takeaway,
                evidence_ids=evidence_ids,
                source_note=source_note,
                steps=tuple(milestones),
            )
        )
    return tuple(exhibits)


def _add_exhibit_title(
    doc: Document,
    exhibit: _RenderableExhibit,
    *,
    compact: bool,
) -> None:
    tag = doc.add_paragraph()
    tag.paragraph_format.space_before = Pt(8)
    tag.paragraph_format.space_after = Pt(4)
    tag.paragraph_format.line_spacing = Pt(10)
    tag.paragraph_format.keep_with_next = True
    run = tag.add_run(exhibit.exhibit_type.upper())
    run.bold = True
    run.font.name = BODY_FONT
    run.font.size = Pt(8)
    run.font.color.rgb = TERMINAL_BLUE

    heading = doc.add_heading(exhibit.title, level=2)
    heading.paragraph_format.space_before = Pt(2)
    heading.paragraph_format.space_after = Pt(5)
    takeaway = doc.add_paragraph()
    takeaway.paragraph_format.space_after = Pt(8)
    takeaway.paragraph_format.keep_with_next = True
    _add_inline(
        takeaway,
        exhibit.takeaway,
        base_size=9.5 if compact else 10.5,
        font=DISPLAY_FONT,
    )
    for run in takeaway.runs:
        run.font.color.rgb = RUNWAY_NAVY
    _style_callout(takeaway)


def _add_tabular_exhibit(
    doc: Document,
    exhibit: _RenderableExhibit,
    *,
    compact: bool,
) -> None:
    table = doc.add_table(rows=1, cols=1 + len(exhibit.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    _set_repeat_table_header(table.rows[0])
    header_labels = [exhibit.row_header] + [
        (
            f"{column.label} ({column.unit})"
            if column.unit
            else column.label
        )
        for column in exhibit.columns
    ]
    for cell, label in zip(table.rows[0].cells, header_labels):
        _set_cell_shading(cell, "0B2D4D")
        _set_cell_margins(cell)
        paragraph = cell.paragraphs[0]
        _add_inline(
            paragraph,
            label,
            base_size=8.25 if compact else 8.75,
            font=BODY_FONT,
        )
        for run in paragraph.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for row_index, exhibit_row in enumerate(exhibit.rows):
        row = table.add_row()
        _prevent_row_split(row)
        values = [exhibit_row.label] + list(exhibit_row.values)
        for column_index, (cell, value) in enumerate(zip(row.cells, values)):
            _set_cell_margins(cell, top=70, start=75, bottom=70, end=75)
            if row_index % 2:
                _set_cell_shading(cell, APRON_FOG)
            if column_index == 0:
                _set_cell_shading(cell, "DDE9EF")
            paragraph = cell.paragraphs[0]
            display = value
            if column_index == 0 and exhibit_row.unit:
                display = f"{value} ({exhibit_row.unit})"
            _add_inline(
                paragraph,
                display,
                base_size=8.25 if compact else 8.75,
                font=BODY_FONT,
            )
            if column_index == 0:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RUNWAY_NAVY


def _add_flow_exhibit(
    doc: Document,
    exhibit: _RenderableExhibit,
    *,
    compact: bool,
) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(0.75)
    table.columns[1].width = Inches(6.55)
    for step_index, step in enumerate(exhibit.steps, 1):
        if step_index > 1:
            connector = table.add_row()
            connector_cell = connector.cells[0].merge(connector.cells[1])
            connector_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            connector_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            run = connector_cell.paragraphs[0].add_run("↓")
            run.font.name = BODY_FONT
            run.font.size = Pt(12)
            run.font.color.rgb = TERMINAL_BLUE
        row = table.add_row()
        _prevent_row_split(row)
        number_cell, content_cell = row.cells
        number_cell.width = Inches(0.75)
        content_cell.width = Inches(6.55)
        _set_cell_margins(number_cell, top=95, start=70, bottom=95, end=70)
        _set_cell_margins(content_cell, top=95, start=110, bottom=95, end=110)
        _set_cell_shading(number_cell, "0B2D4D")
        if step_index % 2 == 0:
            _set_cell_shading(content_cell, APRON_FOG)
        number_p = number_cell.paragraphs[0]
        number_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = number_p.add_run(str(step_index))
        run.bold = True
        run.font.name = DISPLAY_FONT
        run.font.size = Pt(15 if not compact else 13)
        run.font.color.rgb = GUIDANCE_GOLD

        content_p = content_cell.paragraphs[0]
        content_p.paragraph_format.space_after = Pt(2)
        _add_inline(
            content_p,
            step.label,
            base_size=9.25 if compact else 10,
            font=BODY_FONT,
        )
        for run in content_p.runs:
            run.bold = True
            run.font.color.rgb = RUNWAY_NAVY
        detail_p = content_cell.add_paragraph()
        detail_p.paragraph_format.space_after = Pt(2)
        _add_inline(
            detail_p,
            step.detail,
            base_size=8.5 if compact else 9,
            font=BODY_FONT,
        )
        for label, value in (("Owner", step.owner), ("Trigger", step.trigger)):
            if not value:
                continue
            meta = content_cell.add_paragraph()
            meta.paragraph_format.space_after = Pt(1)
            label_run = meta.add_run(f"{label}: ")
            label_run.bold = True
            label_run.font.name = BODY_FONT
            label_run.font.size = Pt(8 if compact else 8.5)
            _add_inline(
                meta,
                value,
                base_size=8 if compact else 8.5,
                font=BODY_FONT,
            )


def _add_timeline_exhibit(
    doc: Document,
    exhibit: _RenderableExhibit,
    *,
    compact: bool,
) -> None:
    include_accountability = any(
        step.owner or step.success_measure for step in exhibit.steps
    )
    table = doc.add_table(
        rows=1,
        cols=3 if include_accountability else 2,
    )
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    _set_repeat_table_header(table.rows[0])
    headers = ["When", "Action"]
    if include_accountability:
        headers.append("Accountability / success test")
    for cell, label in zip(table.rows[0].cells, headers):
        _set_cell_shading(cell, "0B2D4D")
        _set_cell_margins(cell)
        _add_inline(
            cell.paragraphs[0],
            label,
            base_size=8.25 if compact else 8.75,
            font=BODY_FONT,
        )
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for step_index, step in enumerate(exhibit.steps):
        row = table.add_row()
        _prevent_row_split(row)
        for cell in row.cells:
            _set_cell_margins(cell, top=75, start=80, bottom=75, end=80)
            if step_index % 2:
                _set_cell_shading(cell, APRON_FOG)
        _set_cell_shading(row.cells[0], "DDE9EF")
        _add_inline(
            row.cells[0].paragraphs[0],
            step.timing,
            base_size=8.5 if compact else 9,
            font=BODY_FONT,
        )
        for run in row.cells[0].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RUNWAY_NAVY
        action_p = row.cells[1].paragraphs[0]
        _add_inline(
            action_p,
            step.label,
            base_size=8.5 if compact else 9,
            font=BODY_FONT,
        )
        for run in action_p.runs:
            run.bold = True
        if step.detail:
            detail_p = row.cells[1].add_paragraph()
            detail_p.paragraph_format.space_after = Pt(0)
            _add_inline(
                detail_p,
                step.detail,
                base_size=8 if compact else 8.5,
                font=BODY_FONT,
            )
        if include_accountability:
            accountability = [
                f"Owner: {step.owner}" if step.owner else "",
                (
                    f"Success test: {step.success_measure}"
                    if step.success_measure
                    else ""
                ),
            ]
            _add_inline(
                row.cells[2].paragraphs[0],
                "\n".join(value for value in accountability if value),
                base_size=8 if compact else 8.5,
                font=BODY_FONT,
            )


def _add_exhibit_source(
    doc: Document,
    exhibit: _RenderableExhibit,
    *,
    compact: bool,
) -> None:
    source = doc.add_paragraph()
    source.paragraph_format.space_before = Pt(3)
    source.paragraph_format.space_after = Pt(8)
    label = source.add_run("Source: ")
    label.bold = True
    label.font.name = BODY_FONT
    label.font.size = Pt(7.75 if compact else 8.25)
    label.font.color.rgb = OPERATIONS_SLATE
    _add_inline(
        source,
        exhibit.source_note,
        base_size=7.75 if compact else 8.25,
        font=BODY_FONT,
    )
    for run in source.runs:
        run.font.color.rgb = OPERATIONS_SLATE


def _add_decision_exhibits(
    doc: Document,
    visual_brief: dict | None,
    *,
    compact: bool = False,
) -> None:
    """Render complete evidence-bound tables, comparisons, flows, and timelines."""
    exhibits = _renderable_exhibits(visual_brief)
    if not exhibits:
        return

    heading = doc.add_heading("Decision exhibits", level=1)
    heading.paragraph_format.space_after = Pt(5)
    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(9)
    _add_inline(
        intro,
        (
            "The exhibits below render supplied, evidence-bound decision data "
            "directly. Source notes identify the basis provided for each "
            "structure."
        ),
        base_size=9.5 if compact else 10,
        font=BODY_FONT,
    )
    for run in intro.runs:
        run.font.color.rgb = OPERATIONS_SLATE

    for exhibit_index, exhibit in enumerate(exhibits):
        if exhibit_index and exhibit.exhibit_type == "timeline":
            doc.add_page_break()
        _add_exhibit_title(doc, exhibit, compact=compact)
        if exhibit.exhibit_type in {"table", "comparison"}:
            _add_tabular_exhibit(doc, exhibit, compact=compact)
        elif exhibit.exhibit_type == "flow":
            _add_flow_exhibit(doc, exhibit, compact=compact)
        else:
            _add_timeline_exhibit(doc, exhibit, compact=compact)
        _add_exhibit_source(doc, exhibit, compact=compact)


def _footnote_evidence_records(
    markdown: str,
) -> tuple[tuple[str, str, str], ...]:
    """Return final-draft note ID, claim context, and source record."""
    definitions: list[tuple[str, str]] = []
    for line in _strip_comments(markdown).splitlines():
        match = FOOTNOTE_DEF_RE.match(line.strip())
        if match:
            definitions.append((match.group(1), match.group(2).strip()))

    units = _content_units(markdown)
    records: list[tuple[str, str, str]] = []
    for note_id, source in definitions:
        marker = f"[^{note_id}]"
        contexts: list[str] = []
        for unit in units:
            if marker not in unit.text:
                continue
            context = _clip_sentences(FOOTNOTE_MARK_RE.sub("", unit.text), 46)
            if context and context not in contexts:
                contexts.append(context)
            if len(contexts) == 2:
                break
        claim_context = " ".join(contexts) if contexts else (
            "Source note retained in the final draft; no inline marker was detected."
        )
        records.append((note_id, claim_context, source))
    return tuple(records)


def _source_appendix_entries(visual_brief: dict | None) -> tuple[str, ...]:
    """Return source-register text while excluding internal lineage fields."""
    if not isinstance(visual_brief, dict):
        return ()
    payload = visual_brief.get("source_appendix")

    def reader_fields(item: dict) -> list[tuple[str, tuple[str, ...]]]:
        fields: list[tuple[str, tuple[str, ...]]] = []
        for key, value in item.items():
            normalized = re.sub(r"[^a-z0-9]+", "_", str(key).lower()).strip("_")
            if (
                normalized.endswith("_id")
                or normalized.endswith("_ids")
                or normalized
                in {
                    "id",
                    "agent",
                    "agent_id",
                    "claim_id",
                    "claim_ids",
                    "evidence_id",
                    "evidence_ids",
                    "source_id",
                    "source_ids",
                    "record_id",
                    "record_ids",
                }
            ):
                continue
            values = _string_values(value)
            if values:
                label = str(key).replace("_", " ").strip().capitalize()
                fields.append((label, values))
        return fields

    if isinstance(payload, str):
        cleaned = _clean_text(payload)
        return (cleaned,) if cleaned else ()
    if isinstance(payload, list):
        entries: list[str] = []
        for item in payload:
            if isinstance(item, dict):
                parts = [
                    f"{label}: {'; '.join(values)}"
                    for label, values in reader_fields(item)
                ]
                if parts:
                    entries.append(" | ".join(parts))
            else:
                entries.extend(_string_values(item))
        return tuple(entries)
    if isinstance(payload, dict):
        entries = []
        for label, values in reader_fields(payload):
            entries.append(f"{label}: {'; '.join(values)}")
        return tuple(entries)
    return ()


def _add_technical_evidence_appendix(
    doc: Document,
    final_draft_md: str,
    visual_brief: dict | None,
) -> None:
    """Add reader-safe claim and exhibit source records."""
    doc.add_heading("Technical appendix: Evidence register", level=1)
    intro = doc.add_paragraph()
    _add_inline(
        intro,
        (
            "This appendix records the numbered source notes retained in the "
            "final report and the source notes that accompany its exhibits. "
            "It supports review of the cited evidence; it does not replace "
            "review of the underlying source material."
        ),
        base_size=10,
        font=BODY_FONT,
    )

    records = _footnote_evidence_records(final_draft_md)
    if records:
        doc.add_heading("Report source notes", level=2)
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        _set_repeat_table_header(table.rows[0])
        for cell, label in zip(
            table.rows[0].cells,
            ("Note", "Claim context", "Source record"),
        ):
            _set_cell_shading(cell, "0B2D4D")
            _set_cell_margins(cell)
            p = cell.paragraphs[0]
            _add_inline(p, label, base_size=9, font=BODY_FONT)
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for index, (note_id, context, source) in enumerate(records, 1):
            row = table.add_row()
            _prevent_row_split(row)
            for cell in row.cells:
                _set_cell_margins(cell)
                if index % 2 == 0:
                    _set_cell_shading(cell, APRON_FOG)
            note_p = row.cells[0].paragraphs[0]
            note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_inline(note_p, note_id, base_size=9, font=BODY_FONT)
            _add_inline(
                row.cells[1].paragraphs[0],
                context,
                base_size=8.5,
                font=BODY_FONT,
            )
            _add_inline(
                row.cells[2].paragraphs[0],
                source,
                base_size=8.5,
                font=BODY_FONT,
            )
    else:
        p = doc.add_paragraph()
        _add_inline(
            p,
            "No numbered source notes were retained in the final draft.",
            base_size=10,
            font=BODY_FONT,
        )

    visuals = _renderable_exhibits(visual_brief)
    if visuals:
        doc.add_heading("Exhibit source notes", level=2)
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        _set_repeat_table_header(table.rows[0])
        for cell, label in zip(
            table.rows[0].cells,
            ("Exhibit", "Decision takeaway", "Source record"),
        ):
            _set_cell_shading(cell, "0B2D4D")
            _set_cell_margins(cell)
            p = cell.paragraphs[0]
            _add_inline(p, label, base_size=8.5, font=BODY_FONT)
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for index, item in enumerate(visuals, 1):
            row = table.add_row()
            _prevent_row_split(row)
            values = (
                item.title,
                item.takeaway,
                item.source_note,
            )
            for cell, value in zip(row.cells, values):
                _set_cell_margins(cell, top=60, start=70, bottom=60, end=70)
                if index % 2 == 0:
                    _set_cell_shading(cell, APRON_FOG)
                _add_inline(
                    cell.paragraphs[0],
                    value,
                    base_size=8.25,
                    font=BODY_FONT,
                )

    source_appendix = _source_appendix_entries(visual_brief)
    if source_appendix:
        doc.add_heading("Additional exhibit sources", level=2)
        for entry in source_appendix:
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, entry, base_size=9, font=BODY_FONT)


def _content_units(markdown: str) -> list[_ContentUnit]:
    """Parse prose and list items while retaining their nearest heading."""
    text = sanitize_reader_markdown(markdown)
    body_lines = [
        line for line in text.splitlines()
        if not FOOTNOTE_DEF_RE.match(line.strip())
    ]
    units: list[_ContentUnit] = []
    section = ""
    paragraph: list[str] = []
    index = 0

    def flush() -> None:
        nonlocal index
        joined = " ".join(paragraph).strip()
        paragraph.clear()
        if joined:
            units.append(_ContentUnit(section, "paragraph", joined, index))
            index += 1

    for raw in body_lines:
        line = raw.strip()
        heading = HEADING_RE.match(line)
        if heading:
            flush()
            section = heading.group(2).strip()
            continue
        bullet = BULLET_RE.match(line) or ORDERED_RE.match(line)
        if bullet:
            flush()
            units.append(_ContentUnit(section, "bullet", bullet.group(1).strip(), index))
            index += 1
            continue
        if not line:
            flush()
            continue
        if HORIZONTAL_RULE_RE.match(line):
            flush()
            continue
        if line.startswith("|") or PIPE_DIVIDER_RE.match(line):
            flush()
            continue
        paragraph.append(line)
    flush()
    return units


def _plain_for_scoring(text: str) -> str:
    text = FOOTNOTE_MARK_RE.sub("", text)
    text = re.sub(r"[*_`>#]", "", text)
    return " ".join(text.split())


def _clip_sentences(text: str, max_words: int) -> str:
    """Keep complete opening sentences up to a firm decision-brief budget."""
    sentences = re.split(r"(?<=[.!?])\s+(?=[A-Z0-9\[])|(?<=\])\s+(?=[A-Z0-9])", text)
    kept: list[str] = []
    words = 0
    for sentence in sentences:
        count = len(_plain_for_scoring(sentence).split())
        if kept and words + count > max_words:
            break
        kept.append(sentence.strip())
        words += count
        if words >= max_words:
            break
    result = " ".join(kept).strip()
    if not result:
        return text.strip()
    return result


def _rank_units(
    units: list[_ContentUnit],
    *,
    section_terms: tuple[str, ...],
    text_terms: tuple[str, ...] = (),
    prefer_bullets: bool = False,
    minimum_score: float = 1.0,
    reward_numeric: bool = True,
) -> list[_ContentUnit]:
    scored: list[tuple[float, _ContentUnit]] = []
    for unit in units:
        section = unit.section.lower()
        plain = _plain_for_scoring(unit.text).lower()
        word_count = len(plain.split())
        if word_count < 8 or word_count > 260:
            continue
        score = 0.0
        score += sum(4.0 for term in section_terms if term in section)
        score += sum(1.3 for term in text_terms if term in plain)
        if reward_numeric and re.search(
            r"(?:\$[\d,.]+|\b\d+(?:\.\d+)?%|\b\d{1,3}(?:,\d{3})+\b)",
            plain,
        ):
            score += 1.2
        if prefer_bullets and unit.kind == "bullet":
            score += 2.5
        if unit.index < 4:
            score += 0.7
        score -= max(0, word_count - 100) / 80
        if score >= minimum_score:
            scored.append((score, unit))
    return [unit for _, unit in sorted(scored, key=lambda item: (-item[0], item[1].index))]


def _unique_extracts(
    candidates: list[_ContentUnit],
    *,
    count: int,
    max_words: int,
    used: set[str],
) -> tuple[str, ...]:
    selected: list[str] = []
    for candidate in candidates:
        text = _clip_sentences(candidate.text, max_words)
        key = re.sub(r"\W+", " ", _plain_for_scoring(text).lower())[:100]
        if not key or key in used:
            continue
        used.add(key)
        selected.append(text)
        if len(selected) == count:
            break
    return tuple(selected)


def _extract_markdown_section(markdown: str, heading_name: str) -> str:
    """Return one H2 section, stopping at the next H1/H2."""
    lines = _strip_comments(markdown).splitlines()
    collecting = False
    selected: list[str] = []
    for line in lines:
        heading = HEADING_RE.match(line.strip())
        if heading and len(heading.group(1)) <= 2:
            if collecting:
                break
            if heading.group(2).strip().lower() == heading_name.lower():
                collecting = True
                continue
        if collecting:
            selected.append(line)
    return "\n".join(selected).strip()


def _split_recommendation_sentences(text: str) -> tuple[str, ...]:
    text = re.sub(
        r"^\s*\*{0,2}(?:the\s+)?recommendation[^.:\n]*[.:]\*{0,2}\s*",
        "",
        text,
        flags=re.IGNORECASE,
    )
    sentences = re.split(r"(?<=[.!?])\s+(?=[A-Z0-9])", text)
    return tuple(
        _clip_sentences(sentence.strip(), 48)
        for sentence in sentences
        if len(_plain_for_scoring(sentence).split()) >= 7
    )[:5]


def _notes_for_brief(markdown: str, selected_text: str) -> tuple[tuple[str, str], ...]:
    note_defs: dict[str, str] = {}
    for line in _strip_comments(markdown).splitlines():
        match = FOOTNOTE_DEF_RE.match(line.strip())
        if match:
            note_defs[match.group(1)] = match.group(2).strip()
    cited: list[str] = []
    for note_id in FOOTNOTE_MARK_RE.findall(selected_text):
        if note_id in note_defs and note_id not in cited:
            cited.append(note_id)
    return tuple((note_id, note_defs[note_id]) for note_id in cited)


def _explicit_decision_brief(markdown: str) -> DecisionBrief | None:
    """Prefer the verified draft's explicit executive-summary contract."""
    section = _extract_markdown_section(markdown, "Executive summary")
    if not section:
        section = _extract_markdown_section(markdown, "Executive decision brief")
    if not section:
        return None

    units = _content_units(section)
    substantive = [
        unit for unit in units
        if len(_plain_for_scoring(unit.text).split()) >= 8
    ]
    if not substantive:
        return None

    recommendation_unit = next(
        (
            unit
            for unit in substantive
            if unit.kind == "paragraph"
            and re.search(
                r"^\s*\*{0,2}(?:the\s+)?recommendation\b",
                unit.text,
                re.IGNORECASE,
            )
        ),
        None,
    )
    if recommendation_unit is None:
        recommendation_unit = next(
            (
                unit
                for unit in substantive
                if unit.kind == "paragraph"
                and re.search(r"\brecommend(?:ation|ed)?\b", unit.text, re.IGNORECASE)
            ),
            None,
        )
    numbered = [unit for unit in substantive if unit.kind == "bullet"]
    if not numbered:
        numbered = [
            unit for unit in substantive
            if recommendation_unit is None or unit.index != recommendation_unit.index
        ]

    recommendation_text = (
        recommendation_unit.text
        if recommendation_unit is not None
        else substantive[-1].text
    )
    recommendations = _split_recommendation_sentences(recommendation_text)
    bottom_line = _clip_sentences(recommendation_text, 75)

    why_now = tuple(_clip_sentences(unit.text, 52) for unit in numbered[:2])
    evidence = tuple(_clip_sentences(unit.text, 50) for unit in numbered[2:6])
    if not evidence:
        evidence = tuple(_clip_sentences(unit.text, 50) for unit in numbered[:4])

    counter_section = _extract_markdown_section(markdown, "The counter-case, honestly presented")
    if not counter_section:
        counter_section = _extract_markdown_section(markdown, "Counter-case")
    counter_units = [
        unit
        for unit in _content_units(counter_section)
        if unit.kind == "paragraph"
        and len(_plain_for_scoring(unit.text).split()) >= 12
        and not re.search(
            r"\b(?:strongest version|strongest case|honestly presented)\b",
            unit.text,
            re.IGNORECASE,
        )
    ]
    risks = tuple(_clip_sentences(unit.text, 48) for unit in counter_units[:2])

    selected_text = " ".join(
        (bottom_line, *why_now, *evidence, *recommendations, *risks)
    )
    return DecisionBrief(
        bottom_line=bottom_line,
        why_now=why_now,
        evidence=evidence,
        recommendations=recommendations,
        risks=risks,
        notes=_notes_for_brief(markdown, selected_text),
    )


def _build_decision_brief(
    markdown: str,
    *,
    visual_brief: dict | None = None,
    decision_context: DecisionContext | dict | None = None,
) -> DecisionBrief:
    """Distill a verified argument by decision relevance, not document order."""

    def finish(brief: DecisionBrief) -> DecisionBrief:
        return _apply_structured_decision_context(
            _apply_visual_decision_context(brief, visual_brief),
            decision_context,
        )

    explicit = _explicit_decision_brief(markdown)
    if explicit is not None:
        return finish(explicit)

    units = _content_units(markdown)
    if not units:
        return finish(
            DecisionBrief(
                bottom_line="No substantive reader-facing argument was available.",
                why_now=(),
                evidence=(),
                recommendations=(),
                risks=(),
            )
        )
    used: set[str] = set()

    bottom_candidates = _rank_units(
        units,
        section_terms=("executive", "bottom line", "thesis", "conclusion"),
        text_terms=("should", "must", "the evidence", "the central", "this report"),
        reward_numeric=False,
    )
    bottom_extracts = _unique_extracts(
        bottom_candidates or units, count=1, max_words=85, used=used
    )
    if not bottom_extracts:
        fallback = next(
            (
                unit.text
                for unit in units
                if _plain_for_scoring(unit.text).strip()
            ),
            "No substantive reader-facing argument was available.",
        )
        bottom_extracts = (_clip_sentences(fallback, 85),)
    bottom_line = bottom_extracts[0]

    why_candidates = _rank_units(
        units,
        section_terms=("why now", "stakes", "context", "problem", "pressure", "current"),
        text_terms=("risk", "cost", "capacity", "growth", "delay", "decision", "now"),
        reward_numeric=False,
    )
    why_now = _unique_extracts(
        why_candidates or units, count=2, max_words=52, used=used
    )

    evidence_candidates = _rank_units(
        units,
        section_terms=("evidence", "finding", "case study", "peer record", "analysis", "data"),
        text_terms=("percent", "million", "billion", "FAA", "airport", "passenger", "cost"),
    )
    evidence = _unique_extracts(
        evidence_candidates or units, count=4, max_words=48, used=used
    )

    recommendation_candidates = _rank_units(
        units,
        section_terms=("recommend", "action", "implication", "roadmap", "priority", "next"),
        text_terms=("should", "must", "begin", "appoint", "create", "adopt", "fund"),
        prefer_bullets=True,
        reward_numeric=False,
    )
    recommendations = _unique_extracts(
        recommendation_candidates, count=5, max_words=42, used=used
    )
    if not recommendations:
        recommendations = _unique_extracts(
            bottom_candidates, count=3, max_words=42, used=used
        )

    risk_candidates = _rank_units(
        units,
        section_terms=("risk", "counter", "caveat", "constraint", "limit", "trade"),
        text_terms=("however", "unless", "risk", "uncertain", "constraint", "could"),
        reward_numeric=False,
    )
    risks = _unique_extracts(
        risk_candidates, count=3, max_words=45, used=used
    )

    selected_text = " ".join(
        (bottom_line, *why_now, *evidence, *recommendations, *risks)
    )

    return finish(
        DecisionBrief(
            bottom_line=bottom_line,
            why_now=why_now,
            evidence=evidence,
            recommendations=recommendations,
            risks=risks,
            notes=_notes_for_brief(markdown, selected_text),
        )
    )


def _add_section_label(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text.upper())
    run.font.name = BODY_FONT
    run.font.size = Pt(8.5)
    run.font.bold = True
    run.font.color.rgb = TERMINAL_BLUE


def _add_decision_brief(doc: Document, brief: DecisionBrief) -> None:
    heading = doc.add_heading("Executive decision brief", level=1)
    heading.paragraph_format.space_after = Pt(8)

    _add_section_label(doc, "Bottom line")
    callout = doc.add_paragraph()
    callout.paragraph_format.space_after = Pt(8)
    _add_inline(callout, brief.bottom_line, base_size=12, font=DISPLAY_FONT)
    for run in callout.runs:
        run.font.color.rgb = RUNWAY_NAVY
    _style_callout(callout)

    _add_section_label(doc, "Decision ownership and execution")
    execution_table = doc.add_table(rows=0, cols=2)
    execution_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    execution_table.style = "Table Grid"
    execution_rows: list[tuple[str, tuple[str, ...]]] = [
        ("Decision owner", (brief.decision_owner,)),
        ("Approval route", (brief.approval_route,)),
        ("First 90-day action", (brief.first_90_day_action,)),
        ("Success measures", brief.success_measures),
    ]
    if brief.time_horizon:
        execution_rows.insert(3, ("Time horizon", (brief.time_horizon,)))
    for row_index, (label, values) in enumerate(execution_rows):
        row = execution_table.add_row()
        _prevent_row_split(row)
        label_cell, value_cell = row.cells
        _set_cell_margins(label_cell)
        _set_cell_margins(value_cell)
        _set_cell_shading(label_cell, "0B2D4D")
        if row_index % 2:
            _set_cell_shading(value_cell, APRON_FOG)
        label_p = label_cell.paragraphs[0]
        _add_inline(label_p, label, base_size=9, font=BODY_FONT)
        for run in label_p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for value_index, value in enumerate(values):
            value_p = (
                value_cell.paragraphs[0]
                if value_index == 0
                else value_cell.add_paragraph()
            )
            if len(values) > 1:
                value_p.style = "List Bullet"
            value_p.paragraph_format.space_after = Pt(2)
            _add_inline(value_p, value, base_size=9.5, font=BODY_FONT)

    def bullets(label: str, items: tuple[str, ...]) -> None:
        if not items:
            return
        _add_section_label(doc, label)
        for item in items:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
            _add_inline(p, item, base_size=10.5, font=BODY_FONT)

    bullets("Why this matters now", brief.why_now)
    bullets("Evidence that should drive the decision", brief.evidence)

    if brief.recommendations:
        _add_section_label(doc, "Recommended action")
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        _set_repeat_table_header(table.rows[0])
        headers = ("Priority", "Action")
        for index, header_text in enumerate(headers):
            cell = table.rows[0].cells[index]
            _set_cell_shading(cell, "0B2D4D")
            _set_cell_margins(cell)
            p = cell.paragraphs[0]
            _add_inline(p, header_text, base_size=9, font=BODY_FONT)
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for priority, action in enumerate(brief.recommendations, 1):
            row = table.add_row()
            _prevent_row_split(row)
            cells = row.cells
            for cell in cells:
                _set_cell_margins(cell)
                if priority % 2 == 0:
                    _set_cell_shading(cell, APRON_FOG)
            p = cells[0].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_inline(p, str(priority), base_size=10, font=BODY_FONT)
            _add_inline(cells[1].paragraphs[0], action, base_size=9.5, font=BODY_FONT)

    bullets("Risks and conditions", brief.risks)

    if brief.notes:
        _add_section_label(doc, "Decision-brief notes")
        notes_table = doc.add_table(rows=0, cols=2)
        notes_table.autofit = True
        for index in range(0, len(brief.notes), 2):
            row = notes_table.add_row()
            _prevent_row_split(row)
            note_items = brief.notes[index:index + 2]
            if len(note_items) == 1:
                row.cells[0].merge(row.cells[1])
            for column, note_item in enumerate(note_items):
                note_id, note = note_item
                cell = row.cells[column]
                _set_cell_margins(cell, top=35, start=55, bottom=45, end=90)
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                mark = p.add_run(note_id)
                mark.font.superscript = True
                mark.font.size = Pt(8)
                p.add_run("  ")
                _add_inline(p, note, base_size=8.5, font=BODY_FONT)
                for run in p.runs:
                    run.font.color.rgb = OPERATIONS_SLATE


def _decision_brief_to_markdown(brief: DecisionBrief) -> str:
    """Compatibility representation used by callers of _distill_for_summary."""
    lines = [
        "# Executive decision brief",
        "",
        "## Bottom line",
        "",
        brief.bottom_line,
        "",
        "## Decision ownership and execution",
        "",
        f"- **Decision owner:** {brief.decision_owner}",
        f"- **Approval route:** {brief.approval_route}",
        f"- **First 90-day action:** {brief.first_90_day_action}",
        *(
            (f"- **Time horizon:** {brief.time_horizon}",)
            if brief.time_horizon
            else ()
        ),
        *(
            f"- **Success measure:** {measure}"
            for measure in brief.success_measures
        ),
        "",
    ]
    sections = (
        ("Why this matters now", brief.why_now),
        ("Evidence that should drive the decision", brief.evidence),
        ("Recommended action", brief.recommendations),
        ("Risks and conditions", brief.risks),
    )
    for heading, items in sections:
        if not items:
            continue
        lines.extend((f"## {heading}", ""))
        lines.extend(f"- {item}" for item in items)
        lines.append("")
    lines.extend(f"[^{note_id}]: {note}" for note_id, note in brief.notes)
    return "\n".join(lines).strip() + "\n"


def _add_compact_document_header(
    doc: Document,
    *,
    title: str,
    artifact_label: str,
    revision_label: str | None = None,
) -> None:
    """Add a compact identity block without turning a short output into a report."""

    if COUNCIL_LOGO.is_file():
        p = doc.add_paragraph()
        p.add_run().add_picture(str(COUNCIL_LOGO), width=Inches(2.35))

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(23)
    run.font.name = DISPLAY_FONT
    run.font.color.rgb = RUNWAY_NAVY

    if revision_label:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        run = p.add_run(revision_label.upper())
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.name = BODY_FONT
        run.font.color.rgb = GUIDANCE_GOLD

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(
        date.today().strftime("%B %Y")
        + "  •  Transform Airports AI Research Council  •  "
        + artifact_label
    )
    run.font.size = Pt(9)
    run.font.name = BODY_FONT
    run.font.color.rgb = OPERATIONS_SLATE

    _add_ai_accountability_notice(doc, compact=True)


def _add_compact_decision_context(
    doc: Document,
    decision_context: DecisionContext | dict | None,
    *,
    heading: str,
) -> None:
    """Render only authoritative structured fields supplied by the run."""

    context = _coerce_decision_context(decision_context)
    rows: list[tuple[str, tuple[str, ...]]] = []
    if context.decision:
        rows.append(("Decision", (context.decision,)))
    if context.decision_owner:
        rows.append(("Decision owner", (context.decision_owner,)))
    if context.approval_path:
        rows.append(("Approval path", (context.approval_path,)))
    if context.first_action:
        rows.append(("First action", (context.first_action,)))
    if context.time_horizon:
        rows.append(("Time horizon", (context.time_horizon,)))
    if context.success_measures:
        rows.append(("Success measures", context.success_measures))
    if not rows:
        return

    _add_section_label(doc, heading)
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_index, (label, values) in enumerate(rows):
        row = table.add_row()
        _prevent_row_split(row)
        label_cell, value_cell = row.cells
        _set_cell_margins(label_cell, top=60, start=90, bottom=60, end=90)
        _set_cell_margins(value_cell, top=60, start=90, bottom=60, end=90)
        _set_cell_shading(label_cell, "0B2D4D")
        if row_index % 2:
            _set_cell_shading(value_cell, APRON_FOG)
        label_p = label_cell.paragraphs[0]
        _add_inline(label_p, label, base_size=8.5, font=BODY_FONT)
        for run in label_p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for value_index, value in enumerate(values):
            value_p = (
                value_cell.paragraphs[0]
                if value_index == 0
                else value_cell.add_paragraph()
            )
            if len(values) > 1:
                value_p.style = "List Bullet"
            value_p.paragraph_format.space_after = Pt(1)
            _add_inline(value_p, value, base_size=9, font=BODY_FONT)


def _build_article(
    title: str,
    final_draft_md: str,
    out_path: Path,
    *,
    revision_label: str | None = None,
) -> None:
    """Build the continuous narrative without report-only front or back matter."""

    assert_quality(lint_markdown(final_draft_md, location="final draft"))
    doc = Document()
    _configure_document(doc)
    _add_cover_page(
        doc,
        title=title,
        subtitle="A long-form argument for airport leaders",
        revision_label=revision_label,
    )
    _markdown_to_docx(doc, final_draft_md, body_size=11)
    doc.save(out_path)


def _build_compact_output(
    title: str,
    final_draft_md: str,
    out_path: Path,
    *,
    artifact_label: str,
    context_heading: str,
    decision_context: DecisionContext | dict | None,
    revision_label: str | None = None,
) -> None:
    """Build a short decision tool with no report-only appendices or TOC."""

    assert_quality(lint_markdown(final_draft_md, location="final draft"))
    doc = Document()
    _configure_document(doc, compact=True)
    _add_compact_document_header(
        doc,
        title=title,
        artifact_label=artifact_label,
        revision_label=revision_label,
    )
    _add_compact_decision_context(
        doc,
        decision_context,
        heading=context_heading,
    )
    _markdown_to_docx(doc, final_draft_md, body_size=10)
    doc.save(out_path)


def _build_one_page_argument_memo_document(
    *,
    title: str,
    final_draft_md: str,
    out_path: Path,
    body_size: float,
    line_spacing: float,
) -> None:
    """Build the restrained memo used by the focused argument workflow."""

    assert_quality(lint_markdown(final_draft_md, location="final draft"))
    doc = Document()
    _configure_document(doc, compact=True)
    _set_default_font(doc, BODY_FONT, body_size)
    section = doc.sections[0]
    section.top_margin = Inches(0.52)
    section.bottom_margin = Inches(0.52)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)
    section.header_distance = Inches(0.18)
    section.footer_distance = Inches(0.18)

    eyebrow = doc.add_paragraph()
    eyebrow.paragraph_format.space_after = Pt(2)
    run = eyebrow.add_run("EXECUTIVE MEMO  /  FOR DECISION")
    run.bold = True
    run.font.name = BODY_FONT
    run.font.size = Pt(8.5)
    run.font.color.rgb = TERMINAL_BLUE

    heading = doc.add_paragraph()
    heading.paragraph_format.space_after = Pt(2)
    run = heading.add_run(title)
    run.bold = True
    run.font.name = DISPLAY_FONT
    run.font.size = Pt(20)
    run.font.color.rgb = RUNWAY_NAVY

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(5)
    run = meta.add_run(
        date.today().strftime("%B %d, %Y")
        + "  •  Transform Airports AI Council"
    )
    run.font.name = BODY_FONT
    run.font.size = Pt(8.5)
    run.font.color.rgb = OPERATIONS_SLATE
    p_pr = meta._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "D4A24C")
    borders.append(bottom)
    p_pr.append(borders)

    accountability = doc.add_paragraph()
    accountability.paragraph_format.space_after = Pt(4)
    run = accountability.add_run(
        "AI-assisted decision support. The accountable human owner must verify "
        "the sources, local fit, and any action."
    )
    run.font.name = BODY_FONT
    run.font.size = Pt(8)
    run.font.color.rgb = OPERATIONS_SLATE

    memo_markdown = re.sub(
        r"\A\s*#\s+[^\n]+\n+",
        "",
        sanitize_reader_markdown(final_draft_md),
        count=1,
    )
    memo_markdown = _compact_argument_memo_source_urls(memo_markdown)
    body_start = len(doc.paragraphs)
    _markdown_to_docx(
        doc,
        memo_markdown,
        body_size=body_size,
        font=BODY_FONT,
    )
    sources_started = False
    for paragraph in doc.paragraphs[body_start:]:
        style_name = str(getattr(paragraph.style, "name", "") or "")
        if style_name.startswith("Heading"):
            size = 11.5 if style_name == "Heading 2" else 10.5
            if paragraph.text.strip() == "Notes":
                paragraph.runs[0].text = "Sources"
                size = 9
                sources_started = True
            for heading_run in paragraph.runs:
                heading_run.font.size = Pt(size)
            paragraph.paragraph_format.space_before = Pt(3 if sources_started else 5)
            paragraph.paragraph_format.space_after = Pt(1 if sources_started else 2)
            paragraph.paragraph_format.keep_with_next = True
        elif sources_started:
            paragraph.paragraph_format.line_spacing = 0.92
            paragraph.paragraph_format.space_after = Pt(0.5)
            for source_run in paragraph.runs:
                source_run.font.size = Pt(6.75)
        else:
            paragraph.paragraph_format.line_spacing = line_spacing
            paragraph.paragraph_format.space_after = Pt(3.5)

    footer = section.footer
    footer_p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_p.text = ""
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer_p.add_run("HUMAN REVIEW REQUIRED  •  ")
    run.font.name = BODY_FONT
    run.font.size = Pt(8)
    run.font.color.rgb = OPERATIONS_SLATE
    _add_page_field(footer_p)
    doc.save(out_path)


def build_one_page_argument_memo(
    *,
    slug: str,
    title: str,
    final_draft: Path,
    out_dir: Path,
) -> tuple[Path, Path]:
    """Build, render, and page-bind an exact one-page executive memo."""

    out_dir.mkdir(parents=True, exist_ok=True)
    final_draft_md = final_draft.read_text(encoding="utf-8")
    memo_path = out_dir / f"argument-{slug}-memo.docx"
    render_dir = out_dir / "qa" / memo_path.stem
    rendered: list[Path] = []
    render_issues: list[QualityIssue] = []
    page_count = 0
    for body_size, line_spacing in ((10.25, 1.08), (9.5, 1.04)):
        _build_one_page_argument_memo_document(
            title=title,
            final_draft_md=final_draft_md,
            out_path=memo_path,
            body_size=body_size,
            line_spacing=line_spacing,
        )
        rendered, render_issues = render_office_artifact(
            memo_path,
            render_dir,
            required=True,
        )
        page_count = len(
            [path for path in rendered if path.suffix.lower() == ".png"]
        )
        if page_count == 1 or any(issue.severity == "error" for issue in render_issues):
            break

    report = qa_docx(memo_path)
    report.issues.extend(render_issues)
    report.rendered_files.extend(str(path) for path in rendered)
    if page_count != 1:
        report.issues.append(
            QualityIssue(
                code="memo_page_count",
                severity="error",
                message=f"Strengthened-argument memo must render to exactly one page; got {page_count}.",
                location=str(memo_path),
            )
        )

    receipt_path = out_dir / f"{memo_path.stem}-word-visual-inspection.json"
    if report.ok:
        prepare_word_visual_inspection_receipt(
            artifact=memo_path,
            rendered_files=rendered,
            receipt_path=receipt_path,
        )
    bundle = QualityReport(
        artifact=str(memo_path),
        kind="argument_memo_bundle",
        issues=list(report.issues),
        metadata={
            "output_format": "one_page_memo",
            "page_count": page_count,
            "visual_inspection_receipt": str(receipt_path),
        },
        rendered_files=list(report.rendered_files),
    )
    bundle.write_json(out_dir.parent / "publishing-quality.json")
    assert_quality(bundle)
    return memo_path, receipt_path


def _build_brief(
    title: str,
    final_draft_md: str,
    out_path: Path,
    *,
    decision_context: DecisionContext | dict | None,
    revision_label: str | None = None,
) -> None:
    _build_compact_output(
        title,
        final_draft_md,
        out_path,
        artifact_label="Executive brief",
        context_heading="Decision frame",
        decision_context=decision_context,
        revision_label=revision_label,
    )


def _build_recommendations(
    title: str,
    final_draft_md: str,
    out_path: Path,
    *,
    decision_context: DecisionContext | dict | None,
    revision_label: str | None = None,
) -> None:
    _build_compact_output(
        title,
        final_draft_md,
        out_path,
        artifact_label="Action recommendations",
        context_heading="Decision mandate",
        decision_context=decision_context,
        revision_label=revision_label,
    )


def _build_executive_summary(
    title: str,
    final_draft_md: str,
    out_path: Path,
    visual_brief: dict | None = None,
    decision_context: DecisionContext | dict | None = None,
    revision_label: str | None = None,
) -> None:
    """Build a concise decision instrument from the verified argument."""
    assert_quality(lint_markdown(final_draft_md, location="final draft"))
    doc = Document()
    _configure_document(doc, compact=True)

    if COUNCIL_LOGO.is_file():
        p = doc.add_paragraph()
        p.add_run().add_picture(str(COUNCIL_LOGO), width=Inches(2.35))

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(23)
    run.font.name = DISPLAY_FONT
    run.font.color.rgb = RUNWAY_NAVY

    if revision_label:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        run = p.add_run(revision_label.upper())
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.name = BODY_FONT
        run.font.color.rgb = GUIDANCE_GOLD

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(
        date.today().strftime("%B %Y")
        + "  •  Transform Airports AI Research Council  •  Decision brief"
    )
    run.font.size = Pt(9)
    run.font.name = BODY_FONT
    run.font.color.rgb = OPERATIONS_SLATE

    _add_ai_accountability_notice(doc, compact=True)
    _add_decision_brief(
        doc,
        _build_decision_brief(
            final_draft_md,
            visual_brief=visual_brief,
            decision_context=decision_context,
        ),
    )
    doc.save(out_path)


def _distill_for_summary(markdown: str, citation_re: re.Pattern | None = None) -> str:
    """Return a decision-focused summary; retained for API compatibility."""
    text = citation_re.sub("", markdown) if citation_re is not None else markdown
    return _decision_brief_to_markdown(_build_decision_brief(text))


def build_documents(
    *,
    slug: str,
    title: str,
    final_draft: Path,
    methodology: Path,
    out_dir: Path,
    output_format: str = "report",
    visual_brief: Path | None = None,
    decision_context: DecisionContext | dict | None = None,
    revision_label: str | None = None,
) -> tuple[Path, Path | None]:
    """Build format-specific Stage 4 Word artifacts from the verified draft."""
    out_dir.mkdir(parents=True, exist_ok=True)
    final_draft_md = final_draft.read_text(encoding="utf-8")
    requested_format = str(output_format or "").strip().lower()
    if requested_format not in {
        "report",
        "article",
        "brief",
        "recommendations",
    }:
        raise ValueError(f"Unsupported Word output format: {output_format!r}")

    methodology_md = (
        methodology.read_text(encoding="utf-8")
        if requested_format == "report" and methodology.is_file()
        else ""
    )
    visual_payload: dict | None = None
    if visual_brief is not None and visual_brief.is_file():
        payload = json.loads(visual_brief.read_text(encoding="utf-8"))
        if not isinstance(payload, dict):
            raise ValueError(f"Visual brief must contain a JSON object: {visual_brief}")
        visual_payload = payload

    primary_path = out_dir / f"{slug}.docx"
    exec_path: Path | None = None
    if requested_format == "report":
        _build_full_report(
            title,
            final_draft_md,
            methodology_md,
            primary_path,
            visual_brief=visual_payload,
            decision_context=decision_context,
            revision_label=revision_label,
        )
        exec_path = out_dir / f"{slug}-executive-summary.docx"
        _build_executive_summary(
            title,
            final_draft_md,
            exec_path,
            visual_brief=visual_payload,
            decision_context=decision_context,
            revision_label=revision_label,
        )
    elif requested_format == "article":
        _build_article(
            title,
            final_draft_md,
            primary_path,
            revision_label=revision_label,
        )
        exec_path = out_dir / f"{slug}-executive-summary.docx"
        _build_executive_summary(
            title,
            final_draft_md,
            exec_path,
            visual_brief=visual_payload,
            decision_context=decision_context,
            revision_label=revision_label,
        )
    elif requested_format == "brief":
        _build_brief(
            title,
            final_draft_md,
            primary_path,
            decision_context=decision_context,
            revision_label=revision_label,
        )
    else:
        _build_recommendations(
            title,
            final_draft_md,
            primary_path,
            decision_context=decision_context,
            revision_label=revision_label,
        )

    artifacts = [primary_path] + ([exec_path] if exec_path is not None else [])
    reports: list[QualityReport] = []
    visual_inspection_receipts: list[str] = []
    for artifact_path in artifacts:
        report = qa_docx(artifact_path)
        rendered, render_issues = render_office_artifact(
            artifact_path,
            out_dir / "qa" / artifact_path.stem,
            required=True,
        )
        report.issues.extend(render_issues)
        report.rendered_files.extend(str(path) for path in rendered)
        receipt_path = out_dir / (
            f"{artifact_path.stem}-word-visual-inspection.json"
        )
        prepare_word_visual_inspection_receipt(
            artifact=artifact_path,
            rendered_files=rendered,
            receipt_path=receipt_path,
        )
        visual_inspection_receipts.append(str(receipt_path))
        reports.append(report)

    bundle = QualityReport(
        artifact=str(out_dir),
        kind="publishing_bundle",
        issues=[issue for report in reports for issue in report.issues],
        metadata={
            "artifacts": [report.to_dict() for report in reports],
            "output_format": requested_format,
            "visual_brief": str(visual_brief) if visual_brief is not None else None,
            "render_attempted": True,
            "word_visual_inspection_receipts": visual_inspection_receipts,
        },
        rendered_files=[
            rendered
            for report in reports
            for rendered in report.rendered_files
        ],
    )
    quality_path = out_dir.parent / "publishing-quality.json"
    bundle.write_json(quality_path)
    assert_quality(bundle)

    return primary_path, exec_path
