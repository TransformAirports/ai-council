"""Deterministic quality gates for reader-facing Council artifacts.

The Council's model-based review is intentionally complemented by these
mechanical checks.  They require no network access and catch the defects that
language models are poor at policing consistently: leaked internal labels,
unresolved placeholders, broken footnotes, malformed Word files, and missing
document structure.
"""
from __future__ import annotations

import hashlib
import json
import math
import re
import shutil
import subprocess
import tempfile
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from PIL import Image, ImageDraw
from pypdf import PdfReader


PLACEHOLDER_PATTERNS: tuple[tuple[str, re.Pattern[str]], ...] = (
    ("template_token", re.compile(r"\{\{[^{}\n]+\}\}")),
    ("angle_token", re.compile(r"<<[^<>\n]+>>")),
    ("unfinished_marker", re.compile(r"\b(?:TBD|TODO|FIXME)\b", re.IGNORECASE)),
    ("insert_instruction", re.compile(r"\[(?:INSERT|ADD|REPLACE)[^\]]*\]", re.IGNORECASE)),
    ("lorem_ipsum", re.compile(r"\blorem ipsum\b", re.IGNORECASE)),
)

# These are implementation artifacts, not human-readable source citations.
INTERNAL_PATTERNS: tuple[tuple[str, re.Pattern[str]], ...] = (
    (
        "unverified_release_tag",
        re.compile(r"\[UNVERIFIED[^\]]*\]", re.IGNORECASE),
    ),
    (
        "legacy_source_tag",
        re.compile(r"\[Source:\s*[^\]]+\]", re.IGNORECASE),
    ),
    (
        "brief_reference",
        re.compile(
            r"\[[^\]]*(?:\bbrief\b|[-_]brief(?:\.md)?)[^\]]*\]",
            re.IGNORECASE,
        ),
    ),
    (
        "internal_filename",
        re.compile(
            r"\b(?:humanized-draft|edited-draft|final-draft|"
            r"strategist-v\d+|red-team-v\d+)(?:\.md)?\b",
            re.IGNORECASE,
        ),
    ),
    (
        "stage_header",
        re.compile(
            r"(?im)^\s*#{0,6}\s*.*\bStage\s+[123]\b.*"
            r"\b(?:draft|fact[- ]?checked|output)\b.*$"
        ),
    ),
)

FOOTNOTE_MARK_RE = re.compile(r"\[\^([A-Za-z0-9_-]+)\](?!:)")
FOOTNOTE_DEF_RE = re.compile(r"(?m)^\[\^([A-Za-z0-9_-]+)\]:")
MARKDOWN_HEADING_RE = re.compile(r"(?m)^(#{2,6})\s+(.+?)\s*$")
MARKDOWN_LIST_ITEM_RE = re.compile(r"^\s*(?:[-*+]\s+|\d+[.)]\s+)")
MARKDOWN_WORD_RE = re.compile(r"\b[\w]+(?:[-’'][\w]+)*\b", re.UNICODE)
RUN_LENGTH_SECTION_RE = re.compile(
    r"(?ms)^##\s+Length(?:\s+\([^\n)]*\))?\s*$\n"
    r"(?P<body>.*?)(?=^##\s+|\Z)"
)

# These labels describe the writing process instead of the subject. The first
# group is the legacy report scaffold and receives a stronger diagnostic; the
# checks remain advisory because prose-style defects should not strand an
# otherwise verified paid run at the packaging step.
REJECTED_REPORT_HEADINGS = {
    "the argument",
    "why the counter case is insufficient",
    "implications for the operator",
}
GENERIC_SECTION_HEADINGS = REJECTED_REPORT_HEADINGS | {
    "the case",
    "the counter case",
    "the counter case honestly presented",
    "analysis",
    "background",
    "conclusion",
    "discussion",
    "findings",
    "next steps",
    "overview",
}
LIST_FORWARD_SECTION_HEADINGS = {
    "executive summary",
    "decision and bottom line",
    "three findings that carry the decision",
    "recommended action and guardrails",
    "recommendations",
    "sources",
    "source appendix",
    "notes",
    "methodology",
}


@dataclass(frozen=True)
class QualityIssue:
    code: str
    severity: str
    message: str
    location: str = ""


@dataclass
class QualityReport:
    artifact: str
    kind: str
    issues: list[QualityIssue] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)
    rendered_files: list[str] = field(default_factory=list)

    @property
    def errors(self) -> list[QualityIssue]:
        return [issue for issue in self.issues if issue.severity == "error"]

    @property
    def warnings(self) -> list[QualityIssue]:
        return [issue for issue in self.issues if issue.severity == "warning"]

    @property
    def ok(self) -> bool:
        return not self.errors

    def to_dict(self) -> dict[str, Any]:
        return {
            "artifact": self.artifact,
            "kind": self.kind,
            "ok": self.ok,
            "error_count": len(self.errors),
            "warning_count": len(self.warnings),
            "issues": [asdict(issue) for issue in self.issues],
            "metadata": self.metadata,
            "rendered_files": self.rendered_files,
        }

    def write_json(self, path: Path) -> Path:
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(
            json.dumps(self.to_dict(), indent=2, ensure_ascii=False) + "\n",
            encoding="utf-8",
        )
        return path


def lint_reader_text(text: str, *, location: str = "text") -> list[QualityIssue]:
    """Return hard publishing defects found in reader-facing text."""
    issues: list[QualityIssue] = []
    for code, pattern in PLACEHOLDER_PATTERNS:
        for match in pattern.finditer(text):
            excerpt = " ".join(match.group(0).split())[:100]
            issues.append(
                QualityIssue(
                    code=code,
                    severity="error",
                    message=f"Unresolved publishing placeholder: {excerpt!r}",
                    location=location,
                )
            )
    for code, pattern in INTERNAL_PATTERNS:
        for match in pattern.finditer(text):
            excerpt = " ".join(match.group(0).split())[:140]
            issues.append(
                QualityIssue(
                    code=code,
                    severity="error",
                    message=f"Internal Council machinery leaked to the reader: {excerpt!r}",
                    location=location,
                )
            )
    return issues


def _markdown_words(text: str) -> list[str]:
    """Return reader-facing words for conservative prose-density checks."""

    text = re.sub(r"```.*?```", " ", text, flags=re.DOTALL)
    text = re.sub(r"`[^`]+`", " ", text)
    text = re.sub(r"!?(?:\[([^\]]*)\])\([^)]*\)", r"\1", text)
    text = FOOTNOTE_MARK_RE.sub(" ", text)
    return MARKDOWN_WORD_RE.findall(text)


def _normalise_markdown_heading(title: str) -> str:
    title = re.sub(r"[*_`~]", "", title)
    title = re.sub(r"\s+#+\s*$", "", title)
    return " ".join(re.findall(r"[a-z0-9]+", title.casefold()))


def executive_summary_word_target(run_prompt: str) -> int | None:
    """Return an explicit executive-summary target without reading report length.

    The match is intentionally anchored to the words ``executive summary`` so
    a total-report range such as ``4,000–6,000 words`` cannot be mistaken for
    the summary contract that follows it.
    """

    section = RUN_LENGTH_SECTION_RE.search(run_prompt)
    if section is None:
        return None
    length_text = re.sub(r"[*_`~]", "", section.group("body"))
    for clause in re.split(r"[;\n.]", length_text):
        if re.search(r"(?i)\bexecutive\s+summary\b", clause) is None:
            continue
        candidates: list[int] = []
        for raw in re.findall(r"(?<![\w])\d[\d,]*", clause):
            try:
                value = int(raw.replace(",", ""))
            except ValueError:
                continue
            if 100 <= value <= 5_000:
                candidates.append(value)
        if candidates:
            # A range names its upper target last (900–1,100); choosing the
            # maximum also ignores nearby claim counts such as "five claims".
            return max(candidates)
    return None


def _markdown_sections(
    markdown: str,
) -> list[tuple[int, str, str, str]]:
    """Return ``(line, title, normalised title, body)`` for H2-H6 sections."""

    matches = list(MARKDOWN_HEADING_RE.finditer(markdown))
    sections: list[tuple[int, str, str, str]] = []
    for index, match in enumerate(matches):
        level = len(match.group(1))
        body_end = len(markdown)
        for following in matches[index + 1:]:
            if len(following.group(1)) <= level:
                body_end = following.start()
                break
        title = match.group(2).strip()
        sections.append(
            (
                markdown.count("\n", 0, match.start()) + 1,
                title,
                _normalise_markdown_heading(title),
                markdown[match.end():body_end].strip(),
            )
        )
    return sections


def _markdown_paragraphs(markdown: str) -> list[tuple[int, str]]:
    """Return plain prose blocks while excluding headings, lists, and notes."""

    paragraphs: list[tuple[int, str]] = []
    block: list[str] = []
    block_line = 1
    in_fence = False

    def flush() -> None:
        nonlocal block
        text = "\n".join(block).strip()
        block = []
        if not text:
            return
        first = text.lstrip()
        if (
            first.startswith("#")
            or first.startswith("|")
            or first.startswith(">")
            or first.startswith("[^")
            or MARKDOWN_LIST_ITEM_RE.match(first)
        ):
            return
        paragraphs.append((block_line, text))

    for line_number, line in enumerate(markdown.splitlines(), 1):
        if line.lstrip().startswith("```"):
            flush()
            in_fence = not in_fence
            continue
        if in_fence:
            continue
        if not line.strip():
            flush()
            continue
        if not block:
            block_line = line_number
        block.append(line)
    flush()
    return paragraphs


def _lint_narrative_markdown(
    markdown: str,
    *,
    location: str,
    executive_summary_target_words: int = 600,
) -> tuple[list[QualityIssue], dict[str, Any]]:
    """Check prose architecture without pretending to judge literary taste."""

    issues: list[QualityIssue] = []
    sections = _markdown_sections(markdown)
    generic_sections: list[tuple[int, str, str]] = []
    headings_seen: dict[str, tuple[int, str]] = {}
    bodies_seen: dict[str, tuple[int, str]] = {}
    long_paragraphs = 0
    list_heavy_sections = 0

    for line, title, normalised, body in sections:
        if normalised in GENERIC_SECTION_HEADINGS:
            generic_sections.append((line, title, normalised))
            issues.append(
                QualityIssue(
                    code="generic_section_heading",
                    severity="warning",
                    message=(
                        f"Generic heading {title!r} describes report scaffolding; "
                        "replace it with a content-specific assertion."
                    ),
                    location=f"{location}:line {line}",
                )
            )

        if normalised:
            if normalised in headings_seen:
                prior_line, prior_title = headings_seen[normalised]
                issues.append(
                    QualityIssue(
                        code="duplicate_section_heading",
                        severity="warning",
                        message=(
                            f"Heading {title!r} repeats the section at line "
                            f"{prior_line} ({prior_title!r})."
                        ),
                        location=f"{location}:line {line}",
                    )
                )
            else:
                headings_seen[normalised] = (line, title)

        body_without_notes = "\n".join(
            body_line
            for body_line in body.splitlines()
            if not body_line.lstrip().startswith("[^")
        )
        body_words = _markdown_words(body_without_notes)
        if len(body_words) >= 50:
            normalised_body = " ".join(word.casefold() for word in body_words)
            if normalised_body in bodies_seen:
                prior_line, prior_title = bodies_seen[normalised_body]
                issues.append(
                    QualityIssue(
                        code="duplicate_section_body",
                        severity="error",
                        message=(
                            f"Section {title!r} duplicates the reader-facing "
                            f"body of {prior_title!r} at line {prior_line}."
                        ),
                        location=f"{location}:line {line}",
                    )
                )
            else:
                bodies_seen[normalised_body] = (line, title)

        list_lines = [
            body_line
            for body_line in body_without_notes.splitlines()
            if MARKDOWN_LIST_ITEM_RE.match(body_line)
        ]
        list_word_count = sum(len(_markdown_words(item)) for item in list_lines)
        if (
            normalised not in LIST_FORWARD_SECTION_HEADINGS
            and len(list_lines) >= 6
            and len(body_words) >= 120
            and list_word_count / max(1, len(body_words)) >= 0.55
        ):
            list_heavy_sections += 1
            issues.append(
                QualityIssue(
                    code="list_heavy_prose",
                    severity="warning",
                    message=(
                        f"Section {title!r} puts most of its argument in "
                        "list items; convert analytical scaffolding to prose."
                    ),
                    location=f"{location}:line {line}",
                )
            )

        if normalised == "executive summary":
            summary_words = len(body_words)
            summary_target = max(100, int(executive_summary_target_words))
            if summary_words > summary_target:
                issues.append(
                    QualityIssue(
                        code="oversized_executive_summary",
                        severity="warning",
                        message=(
                            f"Executive summary contains {summary_words} words; "
                            f"this run's layered summary target is {summary_target}."
                        ),
                        location=f"{location}:line {line}",
                    )
                )
            numbered_claims = sum(
                1
                for body_line in body.splitlines()
                if re.match(r"^\s*\d+[.)]\s+", body_line)
            )
            if numbered_claims > 5:
                issues.append(
                    QualityIssue(
                        code="too_many_summary_claims",
                        severity="warning",
                        message=(
                            f"Executive summary has {numbered_claims} numbered "
                            "claims; keep the decision layer to four or five."
                        ),
                        location=f"{location}:line {line}",
                    )
                )

    if len(generic_sections) >= 2:
        issues.append(
            QualityIssue(
                code="generic_section_scaffold",
                severity="warning",
                message=(
                    "Multiple generic headings expose the report outline "
                    "instead of telling the reader's argument: "
                    + ", ".join(repr(item[1]) for item in generic_sections)
                    + "."
                ),
                location=location,
            )
        )

    for line, paragraph in _markdown_paragraphs(markdown):
        word_count = len(_markdown_words(paragraph))
        if word_count <= 220:
            continue
        long_paragraphs += 1
        issues.append(
            QualityIssue(
                code="oversized_paragraph",
                severity="warning",
                message=(
                    f"Paragraph contains {word_count} words; split it at a "
                    "real turn in the argument."
                ),
                location=f"{location}:line {line}",
            )
        )

    return issues, {
        "sections": len(sections),
        "generic_headings": len(generic_sections),
        "oversized_paragraphs": long_paragraphs,
        "list_heavy_sections": list_heavy_sections,
        "executive_summary_target_words": executive_summary_target_words,
    }


def lint_markdown(
    markdown: str,
    *,
    location: str = "markdown",
    executive_summary_target_words: int | None = None,
) -> QualityReport:
    """Validate reader-facing Markdown before layout."""
    summary_target = executive_summary_target_words or 600
    issues = lint_reader_text(markdown, location=location)
    narrative_issues, narrative_metadata = _lint_narrative_markdown(
        markdown,
        location=location,
        executive_summary_target_words=summary_target,
    )
    issues.extend(narrative_issues)

    markers = set(FOOTNOTE_MARK_RE.findall(markdown))
    definitions = set(FOOTNOTE_DEF_RE.findall(markdown))
    for missing in sorted(markers - definitions):
        issues.append(
            QualityIssue(
                code="missing_footnote_definition",
                severity="error",
                message=f"Footnote marker [^{missing}] has no definition.",
                location=location,
            )
        )
    for orphaned in sorted(definitions - markers):
        issues.append(
            QualityIssue(
                code="orphaned_footnote_definition",
                severity="warning",
                message=f"Footnote definition [^{orphaned}] is never cited.",
                location=location,
            )
        )

    return QualityReport(
        artifact=location,
        kind="markdown",
        issues=issues,
        metadata={
            "characters": len(markdown),
            "footnote_markers": len(markers),
            "footnote_definitions": len(definitions),
            **narrative_metadata,
        },
    )


def _iter_docx_text(doc: Document) -> Iterable[tuple[str, str]]:
    for index, paragraph in enumerate(doc.paragraphs, 1):
        if paragraph.text.strip():
            yield f"paragraph {index}", paragraph.text
    for table_index, table in enumerate(doc.tables, 1):
        for row_index, row in enumerate(table.rows, 1):
            for cell_index, cell in enumerate(row.cells, 1):
                if cell.text.strip():
                    yield (
                        f"table {table_index}, row {row_index}, cell {cell_index}",
                        cell.text,
                    )


def qa_docx(path: Path, *, require_heading: bool = True) -> QualityReport:
    """Run structural and copy-safety checks against a Word document."""
    path = Path(path)
    issues: list[QualityIssue] = []
    try:
        doc = Document(path)
    except Exception as exc:  # noqa: BLE001 - quality report should survive
        return QualityReport(
            artifact=str(path),
            kind="docx",
            issues=[
                QualityIssue(
                    code="unreadable_docx",
                    severity="error",
                    message=f"Word document cannot be reopened: {exc}",
                    location=str(path),
                )
            ],
        )

    text_items = list(_iter_docx_text(doc))
    all_text = "\n".join(text for _, text in text_items)
    for location, text in text_items:
        issues.extend(lint_reader_text(text, location=location))

    if not all_text.strip():
        issues.append(
            QualityIssue(
                code="empty_document",
                severity="error",
                message="Document contains no reader-facing text.",
                location=str(path),
            )
        )

    headings = [
        p for p in doc.paragraphs
        if p.style is not None and p.style.name.lower().startswith("heading")
    ]
    if require_heading and not headings:
        issues.append(
            QualityIssue(
                code="missing_heading_hierarchy",
                severity="warning",
                message="Document has no semantic heading styles.",
                location=str(path),
            )
        )

    empty_cells = 0
    for table in doc.tables:
        for row in table.rows:
            empty_cells += sum(1 for cell in row.cells if not cell.text.strip())
    if empty_cells:
        issues.append(
            QualityIssue(
                code="empty_table_cells",
                severity="warning",
                message=f"Document contains {empty_cells} empty table cell(s).",
                location=str(path),
            )
        )

    tiny_runs = 0
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.text.strip() and run.font.size is not None and run.font.size.pt < 8:
                tiny_runs += 1
    if tiny_runs:
        issues.append(
            QualityIssue(
                code="tiny_text",
                severity="warning",
                message=f"Document contains {tiny_runs} run(s) below 8 pt.",
                location=str(path),
            )
        )

    footer_xml = "\n".join(
        section.footer._element.xml for section in doc.sections
    )
    has_page_number = "PAGE" in footer_xml
    if len(all_text.split()) > 1200 and not has_page_number:
        issues.append(
            QualityIssue(
                code="missing_page_numbers",
                severity="warning",
                message="Long document has no PAGE field in its footer.",
                location=str(path),
            )
        )

    return QualityReport(
        artifact=str(path),
        kind="docx",
        issues=issues,
        metadata={
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "headings": len(headings),
            "words": len(all_text.split()),
            "page_number_field": has_page_number,
        },
    )


def render_office_artifact(
    path: Path,
    out_dir: Path,
    *,
    required: bool = False,
) -> tuple[list[Path], list[QualityIssue]]:
    """Render DOCX/PPTX to PDF and page PNGs when local tools are available.

    Rendering is a QA aid, never a network call. When ``required`` is true,
    missing or failed LibreOffice/Poppler rendering is release-blocking.
    """
    path = Path(path)
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    issues: list[QualityIssue] = []
    # A retry must not inherit pages from a longer prior render. Remove only
    # the files this artifact owns; leave unrelated QA records untouched.
    stale_outputs = [
        out_dir / f"{path.stem}.pdf",
        out_dir / "montage.png",
        *out_dir.glob(f"{path.stem}-*.png"),
    ]
    for stale in stale_outputs:
        if stale.is_file() or stale.is_symlink():
            stale.unlink()

    office = shutil.which("soffice") or shutil.which("libreoffice")
    if office is None:
        return [], [
            QualityIssue(
                code="render_unavailable",
                severity="error" if required else "warning",
                message="LibreOffice is not installed; structural QA ran without visual rendering.",
                location=str(path),
            )
        ]

    with tempfile.TemporaryDirectory(prefix="council-office-") as profile_dir:
        command = [
            office,
            f"-env:UserInstallation=file://{profile_dir}",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(out_dir),
            str(path),
        ]
        try:
            completed = subprocess.run(
                command,
                check=False,
                capture_output=True,
                text=True,
                timeout=120,
            )
        except (OSError, subprocess.TimeoutExpired) as exc:
            return [], [
                QualityIssue(
                    code="render_failed",
                    severity="error" if required else "warning",
                    message=f"LibreOffice rendering failed: {exc}",
                    location=str(path),
                )
            ]
    pdf_path = out_dir / f"{path.stem}.pdf"
    if completed.returncode != 0 or not pdf_path.is_file():
        return [], [
            QualityIssue(
                code="render_failed",
                severity="error" if required else "warning",
                message=(
                    "LibreOffice did not produce a PDF. "
                    f"{completed.stderr.strip()[:300]}"
                ),
                location=str(path),
            )
        ]

    rendered: list[Path] = [pdf_path]
    pdftoppm = shutil.which("pdftoppm")
    if pdftoppm is None:
        issues.append(
            QualityIssue(
                code="png_render_unavailable",
                severity="error" if required else "warning",
                message="Poppler is not installed; PDF rendered but page PNGs were skipped.",
                location=str(path),
            )
        )
        return rendered, issues

    prefix = out_dir / path.stem
    completed = subprocess.run(
        [pdftoppm, "-png", "-r", "144", str(pdf_path), str(prefix)],
        check=False,
        capture_output=True,
        text=True,
        timeout=120,
    )
    if completed.returncode != 0:
        issues.append(
            QualityIssue(
                code="png_render_failed",
                severity="error" if required else "warning",
                message=f"Page PNG rendering failed: {completed.stderr.strip()[:300]}",
                location=str(path),
            )
        )
        return rendered, issues

    pngs = sorted(out_dir.glob(f"{path.stem}-*.png"))
    rendered.extend(pngs)
    if required and not pngs:
        issues.append(
            QualityIssue(
                code="png_render_missing",
                severity="error",
                message="Required page/slide PNG rendering produced no images.",
                location=str(path),
            )
        )
    return rendered, issues


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with Path(path).open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _safe_inspection_file(root: Path, relative: str) -> Path:
    """Resolve one inspection-record path without escape or symlink tricks."""

    if not relative or Path(relative).is_absolute():
        raise ValueError(f"Word inspection path must be relative: {relative!r}")
    root = root.resolve()
    candidate = root / relative
    resolved = candidate.resolve()
    try:
        resolved.relative_to(root)
    except ValueError as exc:
        raise ValueError(
            f"Word inspection path escapes its artifact root: {relative!r}"
        ) from exc
    cursor = candidate
    while cursor != root:
        if cursor.is_symlink():
            raise ValueError(
                f"Word inspection path uses a symlink: {relative!r}"
            )
        cursor = cursor.parent
    if not resolved.is_file():
        raise ValueError(f"Word inspection file is missing: {relative!r}")
    return resolved


def _pdf_page_count(path: Path) -> int:
    """Read a rendered PDF's authoritative page count or fail closed."""

    try:
        reader = PdfReader(str(path), strict=False)
        count = len(reader.pages)
    except Exception as exc:
        raise ValueError(f"Rendered PDF cannot be read: {exc}") from exc
    if count < 1:
        raise ValueError("Rendered PDF has no pages.")
    return count


def build_page_montage(
    rendered_pages: list[Path],
    montage_path: Path,
    *,
    columns: int = 4,
    thumbnail_width: int = 360,
) -> Path:
    """Build a legible page-sequence contact sheet for document inspection."""

    if not rendered_pages:
        raise ValueError("A Word inspection montage needs at least one page.")
    opened = [Image.open(path).convert("RGB") for path in rendered_pages]
    try:
        aspect = max(
            (image.height / max(image.width, 1) for image in opened),
            default=1.3,
        )
        thumbnail_height = max(240, int(thumbnail_width * aspect))
        label_height = 34
        gutter = 18
        rows = math.ceil(len(opened) / columns)
        width = gutter + columns * (thumbnail_width + gutter)
        height = gutter + rows * (thumbnail_height + label_height + gutter)
        montage = Image.new("RGB", (width, height), "#E9EEF3")
        draw = ImageDraw.Draw(montage)
        for index, image in enumerate(opened):
            row, column = divmod(index, columns)
            x = gutter + column * (thumbnail_width + gutter)
            y = gutter + row * (thumbnail_height + label_height + gutter)
            thumbnail = image.copy()
            thumbnail.thumbnail(
                (thumbnail_width, thumbnail_height),
                Image.Resampling.LANCZOS,
            )
            card = Image.new("RGB", (thumbnail_width, thumbnail_height), "white")
            card.paste(
                thumbnail,
                (
                    (thumbnail_width - thumbnail.width) // 2,
                    (thumbnail_height - thumbnail.height) // 2,
                ),
            )
            montage.paste(card, (x, y))
            draw.text(
                (x, y + thumbnail_height + 7),
                f"PAGE {index + 1:02d}",
                fill="#152B45",
            )
        montage_path.parent.mkdir(parents=True, exist_ok=True)
        montage.save(montage_path, format="PNG", optimize=True)
    finally:
        for image in opened:
            image.close()
    return montage_path


def prepare_word_visual_inspection_receipt(
    *,
    artifact: Path,
    rendered_files: Iterable[Path | str],
    receipt_path: Path,
    inspector: str = "art-director",
) -> Path:
    """Create a pending, hash-bound Word page-inspection packet."""

    artifact = Path(artifact).resolve()
    receipt_path = Path(receipt_path).resolve()
    root = receipt_path.parent.resolve()
    try:
        artifact_relative = artifact.relative_to(root)
    except ValueError as exc:
        raise ValueError(
            f"Word inspection artifact must live under {root}: {artifact}"
        ) from exc

    rendered = [Path(item).resolve() for item in rendered_files]
    pdfs = [path for path in rendered if path.suffix.lower() == ".pdf"]
    pages = [path for path in rendered if path.suffix.lower() == ".png"]
    if len(pdfs) != 1 or not pages:
        raise ValueError(
            "Word inspection requires one rendered PDF and at least one page PNG."
        )

    def page_order(path: Path) -> tuple[int, str]:
        match = re.search(r"-(\d+)\.png$", path.name, re.IGNORECASE)
        return (
            int(match.group(1)) if match else 10**9,
            path.name.casefold(),
        )

    pages = sorted(pages, key=page_order)
    pdf_page_count = _pdf_page_count(pdfs[0])
    page_numbers = [page_order(path)[0] for path in pages]
    if (
        len(pages) != pdf_page_count
        or page_numbers != list(range(1, pdf_page_count + 1))
    ):
        raise ValueError(
            "Word inspection requires exactly one sequential PNG for every "
            f"PDF page; PDF has {pdf_page_count} page(s), render inventory "
            f"contains {len(pages)}."
        )
    for path in (*pdfs, *pages):
        try:
            path.relative_to(root)
        except ValueError as exc:
            raise ValueError(
                f"Word inspection render must live under {root}: {path}"
            ) from exc

    montage_path = build_page_montage(
        pages,
        pages[0].parent / "montage.png",
    )
    payload = {
        "schema_version": "1.0",
        "inspection_type": "word_pages",
        "artifact": {
            "path": artifact_relative.as_posix(),
            "sha256": _sha256(artifact),
            "size_bytes": artifact.stat().st_size,
        },
        "pdf": {
            "path": pdfs[0].relative_to(root).as_posix(),
            "sha256": _sha256(pdfs[0]),
            "size_bytes": pdfs[0].stat().st_size,
        },
        "page_count": pdf_page_count,
        "rendered_pages": [
            {
                "page_number": index,
                "path": page.relative_to(root).as_posix(),
                "sha256": _sha256(page),
                "size_bytes": page.stat().st_size,
            }
            for index, page in enumerate(pages, 1)
        ],
        "montage": {
            "path": montage_path.relative_to(root).as_posix(),
            "sha256": _sha256(montage_path),
            "size_bytes": montage_path.stat().st_size,
        },
        "inspection": {
            "inspector": inspector,
            "full_size_each_page_inspected": False,
            "montage_inspected": False,
            "findings_resolved": False,
            "status": "pending",
            "resolved_findings": [],
            "unresolved_findings": [],
        },
    }
    receipt_path.parent.mkdir(parents=True, exist_ok=True)
    temporary = receipt_path.with_name(f".{receipt_path.name}.tmp")
    temporary.write_text(
        json.dumps(payload, indent=2, ensure_ascii=False) + "\n",
        encoding="utf-8",
    )
    temporary.replace(receipt_path)
    return receipt_path


def qa_word_visual_inspection_receipt(
    receipt_path: Path,
    *,
    artifact: Path,
) -> QualityReport:
    """Validate inspection of every page rendered from the exact Word bytes."""

    receipt_path = Path(receipt_path)
    artifact = Path(artifact)
    issues: list[QualityIssue] = []
    try:
        payload = json.loads(receipt_path.read_text(encoding="utf-8"))
        if not isinstance(payload, dict):
            raise ValueError("receipt root is not an object")
    except (OSError, json.JSONDecodeError, ValueError) as exc:
        return QualityReport(
            artifact=str(receipt_path),
            kind="word_visual_inspection",
            issues=[
                QualityIssue(
                    code="invalid_word_inspection_receipt",
                    severity="error",
                    message=f"Word inspection receipt cannot be read: {exc}",
                    location=str(receipt_path),
                )
            ],
        )

    root = receipt_path.parent

    def error(code: str, message: str, location: str = "") -> None:
        issues.append(
            QualityIssue(
                code=code,
                severity="error",
                message=message,
                location=location or str(receipt_path),
            )
        )

    if payload.get("schema_version") != "1.0":
        error("word_inspection_schema", "Word inspection schema must be 1.0.")
    if payload.get("inspection_type") != "word_pages":
        error(
            "word_inspection_type",
            "Word inspection receipt must declare inspection_type 'word_pages'.",
        )

    artifact_record = payload.get("artifact")
    if not isinstance(artifact_record, dict):
        error("word_inspection_binding_missing", "Receipt has no artifact binding.")
    else:
        try:
            bound_artifact = _safe_inspection_file(
                root, str(artifact_record.get("path") or "")
            )
        except ValueError as exc:
            error("word_inspection_binding_invalid", str(exc))
        else:
            if bound_artifact.resolve() != artifact.resolve():
                error(
                    "word_inspection_binding_mismatch",
                    "Receipt artifact path is not the expected Word file.",
                )
            elif artifact_record.get("sha256") != _sha256(artifact):
                error(
                    "word_inspection_hash_mismatch",
                    "Receipt artifact hash does not match the current Word bytes.",
                )

    pdf_record = payload.get("pdf")
    pdf_page_count: int | None = None
    if not isinstance(pdf_record, dict):
        error("word_inspection_pdf_missing", "Receipt has no rendered PDF binding.")
    else:
        try:
            pdf_path = _safe_inspection_file(
                root, str(pdf_record.get("path") or "")
            )
        except ValueError as exc:
            error("word_inspection_pdf_missing", str(exc))
        else:
            if (
                pdf_path.suffix.lower() != ".pdf"
                or pdf_record.get("sha256") != _sha256(pdf_path)
            ):
                error(
                    "word_inspection_pdf_hash_mismatch",
                    "Rendered PDF bytes changed after inspection.",
                )
            try:
                pdf_page_count = _pdf_page_count(pdf_path)
            except ValueError as exc:
                error("word_inspection_pdf_invalid", str(exc))

    page_count = payload.get("page_count")
    rendered_pages = payload.get("rendered_pages")
    if (
        not isinstance(page_count, int)
        or page_count < 1
        or not isinstance(rendered_pages, list)
        or len(rendered_pages) != page_count
    ):
        error(
            "word_inspection_page_inventory",
            "Receipt must bind exactly one full-size PNG to every rendered page.",
        )
        rendered_pages = []
    if (
        isinstance(page_count, int)
        and pdf_page_count is not None
        and page_count != pdf_page_count
    ):
        error(
            "word_inspection_pdf_page_count_mismatch",
            "Receipt page inventory does not cover every page in the bound PDF.",
        )

    seen_numbers: set[int] = set()
    seen_paths: set[Path] = set()
    for record in rendered_pages:
        if not isinstance(record, dict):
            error("word_inspection_page_inventory", "Rendered-page record is invalid.")
            continue
        number = record.get("page_number")
        if not isinstance(number, int) or number in seen_numbers:
            error(
                "word_inspection_page_sequence",
                "Rendered pages need unique integer page numbers.",
            )
        else:
            seen_numbers.add(number)
        try:
            page_path = _safe_inspection_file(
                root, str(record.get("path") or "")
            )
        except ValueError as exc:
            error("word_inspection_page_missing", str(exc))
            continue
        if page_path in seen_paths:
            error(
                "word_inspection_page_inventory",
                "Every rendered-page record must reference a distinct PNG.",
            )
        seen_paths.add(page_path)
        if (
            page_path.suffix.lower() != ".png"
            or record.get("sha256") != _sha256(page_path)
        ):
            error(
                "word_inspection_page_hash_mismatch",
                "Rendered page bytes changed after inspection.",
                str(page_path),
            )
    if seen_numbers and seen_numbers != set(range(1, int(page_count or 0) + 1)):
        error(
            "word_inspection_page_sequence",
            "Rendered-page numbering does not cover the complete document.",
        )

    montage = payload.get("montage")
    if not isinstance(montage, dict):
        error("word_inspection_montage_missing", "Receipt has no montage binding.")
    else:
        try:
            montage_path = _safe_inspection_file(
                root, str(montage.get("path") or "")
            )
        except ValueError as exc:
            error("word_inspection_montage_missing", str(exc))
        else:
            if (
                montage_path.suffix.lower() != ".png"
                or montage.get("sha256") != _sha256(montage_path)
            ):
                error(
                    "word_inspection_montage_hash_mismatch",
                    "Word montage bytes changed after inspection.",
                )

    inspection = payload.get("inspection")
    if not isinstance(inspection, dict):
        error(
            "word_inspection_attestation_missing",
            "Receipt has no Word-page inspection attestation.",
        )
    else:
        if not str(inspection.get("inspector") or "").strip():
            error(
                "word_inspection_identity_missing",
                "Word inspection receipt must name its inspector.",
            )
        for field in (
            "full_size_each_page_inspected",
            "montage_inspected",
            "findings_resolved",
        ):
            if inspection.get(field) is not True:
                error(
                    "word_inspection_incomplete",
                    f"Word inspection attestation {field!r} is not true.",
                )
        if inspection.get("status") != "pass":
            error(
                "word_inspection_not_passed",
                "Word inspection status must be 'pass' before release.",
            )
        unresolved = inspection.get("unresolved_findings")
        if not isinstance(unresolved, list) or unresolved:
            error(
                "word_inspection_findings_open",
                "Word inspection has unresolved findings.",
            )

    return QualityReport(
        artifact=str(receipt_path),
        kind="word_visual_inspection",
        issues=issues,
        metadata={
            "pages": page_count if isinstance(page_count, int) else 0,
            "inspector": (
                inspection.get("inspector")
                if isinstance(inspection, dict)
                else None
            ),
            "receipt_sha256": _sha256(receipt_path),
        },
    )


def assert_quality(report: QualityReport) -> None:
    """Raise a compact error when a deterministic publishing gate fails."""
    if report.ok:
        return
    details = "; ".join(f"{issue.code}: {issue.message}" for issue in report.errors[:8])
    raise ValueError(f"Publishing quality gate failed for {report.artifact}: {details}")
