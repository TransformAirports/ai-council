from __future__ import annotations

import json
import tempfile
import unittest
from pathlib import Path
from unittest.mock import AsyncMock, patch

from fastapi.testclient import TestClient

import cli.server as server
from cli.agents import load_all_agents
from cli.presentation_qa import PresentationQAConfig
from cli.run_manifest import ResumeContractMismatch
from cli.strengthen import (
    StrengthenRequest,
    _prepare_argument_manifest,
    _publish_strengthen_release,
    run_strengthen_pipeline,
)


ALLOWED = {"operations-analyst", "contrarian"}
ORIGIN = "http://127.0.0.1:8723"
CLIENT_ID = "argument_client_1234567890"


class StrengthenRequestTests(unittest.TestCase):
    def test_accepts_pasted_text_without_a_deck(self) -> None:
        request = StrengthenRequest.from_payload(
            {
                "title": "The operating case",
                "argument_text": "The airport should change its operating model.",
                "agents": ["operations-analyst", "operations-analyst"],
                "want_pptx": False,
            },
            allowed_agents=ALLOWED,
        )
        self.assertEqual(request.selected_agents, ["operations-analyst"])
        self.assertIsNone(request.slide_count)

    def test_accepts_documents_only_and_an_exact_slide_count(self) -> None:
        request = StrengthenRequest.from_payload(
            {
                "title": "Board position",
                "source_tokens": ["position.docx"],
                "agents": ["contrarian"],
                "want_pptx": True,
                "slide_count": 7,
            },
            allowed_agents=ALLOWED,
        )
        self.assertEqual(request.source_tokens, ["position.docx"])
        self.assertEqual(request.slide_count, 7)

    def test_rejects_missing_material_agents_and_bad_slide_counts(self) -> None:
        with self.assertRaisesRegex(ValueError, "Paste an argument"):
            StrengthenRequest.from_payload(
                {"title": "Empty", "agents": ["contrarian"]},
                allowed_agents=ALLOWED,
            )
        with self.assertRaisesRegex(ValueError, "Seat at least one"):
            StrengthenRequest.from_payload(
                {"title": "No council", "argument_text": "A claim."},
                allowed_agents=ALLOWED,
            )
        with self.assertRaisesRegex(ValueError, "between 3 and 30"):
            StrengthenRequest.from_payload(
                {
                    "title": "Too long",
                    "argument_text": "A claim.",
                    "agents": ["contrarian"],
                    "want_pptx": True,
                    "slide_count": 31,
                },
                allowed_agents=ALLOWED,
            )
        with self.assertRaisesRegex(ValueError, "whole number"):
            StrengthenRequest.from_payload(
                {
                    "title": "Fractional deck",
                    "argument_text": "A claim.",
                    "agents": ["contrarian"],
                    "want_pptx": True,
                    "slide_count": 7.5,
                },
                allowed_agents=ALLOWED,
            )

    def test_argument_presentation_contract_is_exact(self) -> None:
        config = PresentationQAConfig.for_argument(9)
        self.assertEqual(config.deck_mode, "argument_brief")
        self.assertEqual(config.min_slide_count, 9)
        self.assertEqual(config.max_slide_count, 9)
        self.assertEqual(config.slide_count_severity, "error")
        with self.assertRaisesRegex(ValueError, "between 3 and 30"):
            PresentationQAConfig.for_argument(9.5)  # type: ignore[arg-type]

    def test_resume_identity_rejects_changed_source_bytes(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            outputs = root / "outputs"
            outputs.mkdir()
            source = root / "sources" / "runs" / "operating-case" / "position.md"
            source.parent.mkdir(parents=True)
            source.write_text("The original position.\n", encoding="utf-8")
            request = StrengthenRequest(
                title="Operating case",
                selected_agents=["contrarian"],
                source_tokens=["position.md"],
                slug="operating-case",
            )
            agents = {agent.name: agent for agent in load_all_agents()}
            arguments = {
                "request": request,
                "repo_root": root,
                "outputs_dir": outputs,
                "readable_sources": [source],
                "agents": agents,
                "process_names": ["evidence-curator", "strategist", "fact-checker"],
            }
            _prepare_argument_manifest(**arguments, resume=False)
            _prepare_argument_manifest(**arguments, resume=True)
            source.write_text("The position changed.\n", encoding="utf-8")
            with self.assertRaisesRegex(ResumeContractMismatch, "source bytes"):
                _prepare_argument_manifest(**arguments, resume=True)


class StrengthenReleaseTests(unittest.TestCase):
    def test_release_is_readable_only_while_hashes_match(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            reports = root / "reports"
            archive = root / "runs" / "2026-07-27-argument-operating-case"
            archive.mkdir(parents=True)
            final = root / "final.md"
            memo = root / "memo.docx"
            deck = root / "deck.pptx"
            final.write_text("# Operating case\n\nA verified argument.\n", encoding="utf-8")
            memo.write_bytes(b"validated one-page memo bytes")
            deck.write_bytes(b"validated presentation bytes")
            request = StrengthenRequest(
                title="Operating case",
                argument_text="Initial case.",
                selected_agents=["contrarian"],
                want_pptx=True,
                slide_count=6,
                slug="operating-case",
            )
            argument, word_memo, presentation = _publish_strengthen_release(
                repo_root=root,
                request=request,
                final_argument=final,
                memo=memo,
                deck=deck,
                archive_dir=archive,
            )
            release = server._verified_argument_release(
                "argument-operating-case", reports
            )
            self.assertIsNotNone(release)
            self.assertEqual(argument, reports / "argument-operating-case.md")
            self.assertEqual(
                word_memo, reports / "argument-operating-case-memo.docx"
            )
            self.assertEqual(presentation, reports / "argument-operating-case.pptx")
            argument.write_text("tampered", encoding="utf-8")
            self.assertIsNone(
                server._verified_argument_release("argument-operating-case", reports)
            )


class StrengthenUploadTests(unittest.TestCase):
    def setUp(self) -> None:
        self.client = TestClient(
            server.app,
            base_url=ORIGIN,
            client=("127.0.0.1", 51000),
        )

    def tearDown(self) -> None:
        self.client.close()

    def test_authenticated_upload_and_delete_are_client_scoped(self) -> None:
        with tempfile.TemporaryDirectory() as directory, patch.object(
            server, "REPO_ROOT", Path(directory)
        ):
            headers = {
                "origin": ORIGIN,
                "x-council-session": server._SESSION_TOKEN,
                "x-council-client": CLIENT_ID,
                "content-type": "application/octet-stream",
            }
            response = self.client.post(
                "/api/argument-source?name=position.txt",
                headers=headers,
                content=b"The current argument.",
            )
            self.assertEqual(response.status_code, 200, response.text)
            token = response.json()["token"]
            staged = server._resolve_argument_uploads(
                CLIENT_ID, [token], Path(directory)
            )
            self.assertEqual(staged[0].read_bytes(), b"The current argument.")
            removed = self.client.delete(
                f"/api/argument-source?token={token}", headers=headers
            )
            self.assertEqual(removed.status_code, 200, removed.text)
            self.assertFalse(staged[0].exists())

    def test_report_and_scope_uploads_are_isolated(self) -> None:
        with tempfile.TemporaryDirectory() as directory, patch.object(
            server, "REPO_ROOT", Path(directory)
        ):
            headers = {
                "origin": ORIGIN,
                "x-council-session": server._SESSION_TOKEN,
                "x-council-client": CLIENT_ID,
                "content-type": "application/octet-stream",
            }
            report = self.client.post(
                "/api/source?purpose=report&name=material.txt",
                headers=headers,
                content=b"report material",
            )
            scope = self.client.post(
                "/api/source?purpose=scope&name=material.txt",
                headers=headers,
                content=b"scope material",
            )
            self.assertEqual(report.status_code, 200, report.text)
            self.assertEqual(scope.status_code, 200, scope.text)

            report_path = server._resolve_argument_uploads(
                CLIENT_ID,
                [report.json()["token"]],
                Path(directory),
                purpose="report",
            )[0]
            scope_path = server._resolve_argument_uploads(
                CLIENT_ID,
                [scope.json()["token"]],
                Path(directory),
                purpose="scope",
            )[0]
            self.assertNotEqual(report_path, scope_path)
            self.assertEqual(report_path.read_bytes(), b"report material")
            self.assertEqual(scope_path.read_bytes(), b"scope material")

            rejected = self.client.post(
                "/api/source?purpose=unknown&name=material.txt",
                headers=headers,
                content=b"not accepted",
            )
            self.assertEqual(rejected.status_code, 400, rejected.text)


class StrengthenPipelineTests(unittest.IsolatedAsyncioTestCase):
    async def test_pipeline_releases_the_verified_argument_and_one_page_word_memo(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            (root / "outputs").mkdir()
            request = StrengthenRequest(
                title="A bounded operating change",
                argument_text="The airport should authorize a bounded operating change.",
                research_goal="Test the mechanism and strongest objection.",
                selected_agents=["contrarian"],
            )

            reader_claim = (
                "The official operating record supports a bounded test and limits "
                "it to one terminal."
            )
            evidence_record = {
                "evidence_id": "E-1",
                "claim": reader_claim,
                "source_title": "Official operating record",
                "source_url": "https://example.gov/record",
                "source_type": "official_statement",
                "is_primary": True,
                "page_or_section": "Operating limits",
                "supporting_excerpt": (
                    "A bounded test is authorized and limited to one terminal."
                ),
                "source_date": "2026-01-01",
                "data_vintage": "2026",
                "airport_or_entity": "Test Airport",
                "units": None,
                "denominator": None,
                "caveat": "Pilot only.",
                "confidence": "high",
            }
            filler_paragraphs = [
                " ".join(["Evidence improves the reasoning."] * 22)
                for _ in range(4)
            ]
            filler = "\n\n".join(filler_paragraphs)
            final_text = (
                "# A bounded operating change\n\n"
                "## Bottom line\n\n"
                "Authorize the bounded test and keep the scope narrow.\n\n"
                "## Why it holds\n\n"
                + filler
                + f"\n\n{reader_claim}[^1]\n\n"
                "## Strongest objection\n\n"
                "A narrow test may not prove the model at full scale. That is why the "
                "first decision should buy evidence, not lock in expansion.\n\n"
                "## What to do now\n\n"
                "Name an owner, set the operating limits, and authorize the test.\n\n"
                "[^1]: Official operating record, 2026, Operating limits.\n"
            )
            fact_steps: list[str] = []

            async def fake_run_agent(**kwargs):
                fact_steps.append(str(kwargs.get("step_label") or ""))
                output = Path(kwargs["output_path"])
                output.parent.mkdir(parents=True, exist_ok=True)
                if output.name.endswith("-brief.md"):
                    output.write_text("# Focused brief\n\n" + filler, encoding="utf-8")
                elif output.name == "evidence-map.md":
                    output.write_text("# Argument kit\n\n" + filler, encoding="utf-8")
                elif output.name == "strategist-draft.md":
                    output.write_text(final_text, encoding="utf-8")
                elif output.name == "final-draft.md":
                    output.write_text(final_text, encoding="utf-8")
                for required, _contract in kwargs.get("required_outputs", ()):
                    required = Path(required)
                    required.parent.mkdir(parents=True, exist_ok=True)
                    if required.name.endswith("-evidence.jsonl"):
                        required.write_text(
                            json.dumps(evidence_record) + "\n", encoding="utf-8"
                        )
                    elif required.name == "evidence-map.md":
                        required.write_text(
                            "# Argument kit\n\n" + filler, encoding="utf-8"
                        )
                    elif required.name == "fact-check-report.md":
                        required.write_text(
                            "# Verification\n\n" + "Verified claim. " * 50,
                            encoding="utf-8",
                        )
                    elif required.name == "claim-lineage.jsonl":
                        claim = (
                            reader_claim
                            if kwargs.get("step_label")
                            == "argument/fact-check-remediation"
                            else "The official operating record limits the test to one terminal."
                        )
                        required.write_text(
                            json.dumps(
                                {
                                    "claim_id": "claim-0001",
                                    "claim": claim,
                                    "citation": "Official operating record, 2026, Operating limits.",
                                    "footnote_id": "1",
                                    "evidence_ids": ["contrarian::E-1"],
                                    "verification_status": "verified",
                                    "verification_note": "Primary record checked.",
                                    "primary_source_checked": True,
                                    "retained": True,
                                }
                            )
                            + "\n",
                            encoding="utf-8",
                        )
                return {"skipped": False, "cost": 0.0, "turns": 1, "provider": "anthropic"}

            def fake_build_memo(*, slug, title, final_draft, out_dir):
                from docx import Document

                out_dir.mkdir(parents=True, exist_ok=True)
                memo = out_dir / f"argument-{slug}-memo.docx"
                document = Document()
                document.add_heading(title, level=1)
                document.add_paragraph(final_draft.read_text(encoding="utf-8"))
                document.save(memo)
                receipt = out_dir / f"{memo.stem}-word-visual-inspection.json"
                receipt.write_text("{}\n", encoding="utf-8")
                (out_dir.parent / "publishing-quality.json").write_text(
                    json.dumps(
                        {
                            "artifact": str(memo),
                            "kind": "argument_memo_bundle",
                            "ok": True,
                            "issues": [],
                        }
                    )
                    + "\n",
                    encoding="utf-8",
                )
                return memo, receipt

            with (
                patch("cli.strengthen._run_agent", side_effect=fake_run_agent),
                patch("cli.strengthen.emit", new=AsyncMock()),
                patch("cli.strengthen._notify_done"),
                patch(
                    "cli.docx_builder.build_one_page_argument_memo",
                    side_effect=fake_build_memo,
                ),
                patch(
                    "cli.orchestrator.run_word_visual_inspection",
                    new=AsyncMock(),
                ),
            ):
                result = await run_strengthen_pipeline(
                    request=request, repo_root=root, budget_usd=60
                )

            self.assertTrue(result.completed)
            self.assertTrue(result.argument_path.is_file())
            self.assertTrue(result.memo_path.is_file())
            self.assertIsNone(result.deck_path)
            self.assertEqual(
                [path.name for path in (root / "reports").glob("*.docx")],
                ["argument-a-bounded-operating-change-memo.docx"],
            )
            self.assertIsNotNone(
                server._verified_argument_release(result.public_slug, root / "reports")
            )
            self.assertTrue(result.archive_path.is_dir())
            self.assertIn("argument/fact-check-remediation", fact_steps)


if __name__ == "__main__":
    unittest.main()
