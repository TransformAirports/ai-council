from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from cli.docx_builder import (
    APRON_FOG,
    DECISION_BRIEF_MAX_ITEMS,
    DECISION_BRIEF_MAX_WORDS,
    FOOTNOTE_MARK_RE,
    DecisionBrief,
    _bound_decision_brief,
    _build_executive_summary,
    _build_full_report,
    _clip_sentences,
    _plain_for_scoring,
    _without_front_summary,
)


SUMMARY_CLAIM = (
    "The unmistakable summary claim is that a reversible operating test "
    "protects the morning peak."
)
NARRATIVE_CLAIM = (
    "The main narrative remains intact and explains how the operating team "
    "will pause the test before queues reach the agreed threshold."
)

FINAL_DRAFT = f"""# A reversible terminal operating test

## Executive summary

1. {SUMMARY_CLAIM}[^1]

2. The current queue reaches its limit before seven in the morning, so the
operating window must stay outside that peak.[^2]

3. A daily admission cap keeps the exposure observable and reversible.[^3]

**The recommendation.** Authorize one 90-day off-peak test. Give the duty
manager a pause switch. Review the evidence every two weeks.

## The operating window is the control

{NARRATIVE_CLAIM}[^2]

The test admits visitors only after the morning bank clears. The duty manager
can suspend admission without waiting for a committee meeting.[^3]

## A bounded authorization is enough

Approve the operating envelope, assign one accountable owner, and publish the
pause threshold before launch.[^4]

[^1]: Board operating record, July 2026, https://www.example-airport.test/records/very/long/source/path.
[^2]: Airport queue study, June 2026.
[^3]: Peer-airport pilot record, May 2026.
[^4]: Draft authorization memorandum, August 2026.
"""


def _document_text(path: Path) -> str:
    document = Document(path)
    blocks = [paragraph.text for paragraph in document.paragraphs]
    blocks.extend(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    return "\n".join(blocks)


class DocxReadabilityTests(unittest.TestCase):
    def test_full_package_uses_summary_once_and_preserves_narrative_and_notes(
        self,
    ) -> None:
        transformed = _without_front_summary(FINAL_DRAFT)
        self.assertNotIn(SUMMARY_CLAIM, transformed)
        self.assertIn(NARRATIVE_CLAIM, transformed)
        for note_id in range(1, 5):
            self.assertIn(f"[^{note_id}]:", transformed)

        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            report_path = root / "scratch-report.docx"
            summary_path = root / "scratch-executive-summary.docx"
            _build_full_report(
                "A reversible terminal operating test",
                FINAL_DRAFT,
                "# Methodology\n\nThe Council compared operating records.",
                report_path,
            )
            _build_executive_summary(
                "A reversible terminal operating test",
                FINAL_DRAFT,
                summary_path,
            )

            report_text = _document_text(report_path)
            standalone_text = _document_text(summary_path)
            self.assertEqual(report_text.count(SUMMARY_CLAIM), 1)
            self.assertNotIn("Executive summary", report_text)
            self.assertIn(NARRATIVE_CLAIM, report_text)
            self.assertIn("Board operating record, July 2026", report_text)
            self.assertIn(SUMMARY_CLAIM, standalone_text)

            # Execution labels now scan as quiet metadata instead of a column
            # of oversized dark blocks.
            report = Document(report_path)
            label_cell = report.tables[0].cell(0, 0)
            shading = label_cell._tc.get_or_add_tcPr().find(qn("w:shd"))
            self.assertIsNotNone(shading)
            self.assertEqual(shading.get(qn("w:fill")), APRON_FOG)

    def test_reader_facing_word_files_remove_internal_lineage_tokens(self) -> None:
        marked_up = FINAL_DRAFT.replace(
            SUMMARY_CLAIM,
            SUMMARY_CLAIM + " [operations-analyst::ev-014]",
        ).replace(
            NARRATIVE_CLAIM,
            NARRATIVE_CLAIM + " [strategist::claim-003]",
        ).replace(
            "Airport queue study, June 2026.",
            "Airport queue study, June 2026. [researcher::ev-022]",
        )

        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "scratch-report.docx"
            _build_full_report(
                "A reversible terminal operating test",
                marked_up,
                "# Methodology\n\nThe Council compared operating records.",
                path,
            )
            reader_text = _document_text(path)

        self.assertNotIn("::", reader_text)
        self.assertNotIn("ev-014", reader_text)
        self.assertNotIn("claim-003", reader_text)
        self.assertIn(SUMMARY_CLAIM, reader_text)
        self.assertIn(NARRATIVE_CLAIM, reader_text)

    def test_decision_brief_has_deterministic_content_budgets(self) -> None:
        def long_text(label: str, words: int = 55) -> str:
            sentence = f"{label} remains supported by the verified operating record."
            repeats = max(2, (words // len(sentence.split())) + 1)
            return " ".join(sentence for _ in range(repeats))

        unbounded = DecisionBrief(
            bottom_line=long_text("Authorize"),
            why_now=tuple(long_text(f"Why {index}") for index in range(5)),
            evidence=tuple(long_text(f"Evidence {index}") for index in range(7)),
            recommendations=tuple(
                long_text(f"Recommendation {index}") for index in range(6)
            ),
            risks=tuple(long_text(f"Risk {index}") for index in range(5)),
            notes=tuple(
                (
                    str(index),
                    long_text(f"Source {index}", 30)
                    + f" https://www.source{index}.test/a/very/long/path"
                    + " [research-agent::ev-999]",
                )
                for index in range(1, 7)
            ),
            decision_owner=long_text("Chief Operating Officer"),
            approval_route=long_text("Committee then board"),
            first_90_day_action=long_text("Open one controlled test"),
            success_measures=tuple(
                long_text(f"Measure {index}") for index in range(5)
            ),
            time_horizon=long_text("Ninety days"),
        )
        brief = _bound_decision_brief(unbounded)

        for field in (
            "why_now",
            "evidence",
            "recommendations",
            "risks",
            "success_measures",
        ):
            values = getattr(brief, field)
            self.assertLessEqual(len(values), DECISION_BRIEF_MAX_ITEMS[field])
            self.assertEqual(len(values), len(set(values)))
            for value in values:
                self.assertLessEqual(
                    len(_plain_for_scoring(value).split()),
                    DECISION_BRIEF_MAX_WORDS[field],
                )

        for field in (
            "bottom_line",
            "decision_owner",
            "approval_route",
            "first_90_day_action",
            "time_horizon",
        ):
            self.assertLessEqual(
                len(_plain_for_scoring(getattr(brief, field)).split()),
                DECISION_BRIEF_MAX_WORDS[field],
            )

        # The unbounded fixture has no retained citation markers, so its source
        # records must not leak into the brief merely to fill a notes budget.
        self.assertEqual(brief.notes, ())

    def test_sentence_safe_trimming_never_bisects_a_verified_sentence(self) -> None:
        verified_sentence = (
            "The airport may authorize the test only after the operations team "
            "confirms the queue threshold and the board approves the published "
            "stop condition.[^7]"
        )

        trimmed = _clip_sentences(verified_sentence, max_words=12)

        self.assertEqual(trimmed, "")
        self.assertNotIn("…", trimmed)
        self.assertNotIn("The airport may authorize", trimmed)

    def test_sentence_safe_trimming_can_skip_an_overlong_opening_sentence(self) -> None:
        overlong = (
            "The airport may authorize the test only after the operations team "
            "confirms the queue threshold and the board approves the published "
            "stop condition.[^7]"
        )
        later_complete = "Keep the morning peak outside the test window.[^8]"

        trimmed = _clip_sentences(
            f"{overlong} {later_complete}",
            max_words=12,
        )

        self.assertEqual(trimmed, later_complete)
        self.assertNotIn("[^7]", trimmed)
        self.assertIn("[^8]", trimmed)

    def test_bounding_keeps_retained_markers_and_source_notes_in_lockstep(self) -> None:
        overlong = (
            "This evidence sentence deliberately exceeds the decision brief "
            "budget while preserving a source marker that must disappear with "
            "the complete sentence rather than being separated from it.[^3]"
        )
        brief = DecisionBrief(
            bottom_line="Authorize the reversible operating test.[^1]",
            why_now=(
                "The morning queue establishes the safe operating window.[^2]",
            ),
            evidence=(
                overlong + " A shorter verified finding still fits.[^4]",
            ),
            recommendations=("Publish the pause threshold before launch.[^5]",),
            risks=("Stop the test if the morning queue crosses its limit.[^6]",),
            notes=(
                ("1", "Board authorization, August 2026."),
                ("2", "Queue study, June 2026."),
                ("3", "Long-form operating analysis, July 2026."),
                ("4", "Pilot finding, May 2026."),
                (
                    "5",
                    "Operating directive, August 2026, "
                    "https://www.example-airport.test/directives/pause-threshold "
                    "[research-agent::ev-005]",
                ),
                ("6", "Published operating stop rule, August 2026."),
                ("9", "Unused source record."),
            ),
        )

        bounded = _bound_decision_brief(brief)
        retained_text = " ".join(
            (
                bounded.bottom_line,
                *bounded.why_now,
                *bounded.evidence,
                *bounded.recommendations,
                *bounded.risks,
                bounded.decision_owner,
                bounded.approval_route,
                bounded.first_90_day_action,
                *bounded.success_measures,
                bounded.time_horizon,
            )
        )
        retained_markers = list(dict.fromkeys(FOOTNOTE_MARK_RE.findall(retained_text)))
        retained_notes = [note_id for note_id, _ in bounded.notes]

        self.assertEqual(retained_markers, ["1", "2", "4", "5", "6"])
        self.assertEqual(retained_notes, retained_markers)
        self.assertNotIn("[^3]", retained_text)
        self.assertNotIn("9", retained_notes)
        note_five = dict(bounded.notes)["5"]
        self.assertNotIn("https://", note_five)
        self.assertNotIn("::", note_five)
        self.assertIn("example-airport.test", note_five)


if __name__ == "__main__":
    unittest.main()
